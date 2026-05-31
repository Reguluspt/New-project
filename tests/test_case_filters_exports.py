from __future__ import annotations

import unittest
import tempfile
from pathlib import Path
from unittest.mock import patch
from zipfile import ZipFile

from docx import Document

from src.case_files import case_folder, save_original_file
from src.case_exports import (
    approve_case_documents_pdf,
    collect_template_errors,
    compare_case_documents,
    document_action_error,
    export_case_documents,
    missing_mandatory_data_fields,
    package_case_documents,
)
from src.case_output_preferences import load_case_output_dir, save_case_output_dir
from src.document_exporter import render_docx_preview_html
from src.document_exporter import render_docx_template
from src.document_exporter import build_placeholder_context
from src.case_filters import (
    build_chart_data,
    build_chart_rows,
    build_filters,
    export_scope_label,
    filter_value,
    export_rows_for_filters,
    get_unpaid_report,
    search_page,
)
from src.sqlite_store import CANCELED_CASE_STATUS, DEFAULT_CASE_STATUS, create_case, get_case, init_db, update_case
from src.sqlite_store import distinct_case_values
from src.template_manager import TEMPLATE_REQUIREMENTS, read_docx_text


def _write_docx_template(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph(text)
    document.save(path)


def _write_individual_templates(directory: Path) -> None:
    templates = {
        "mau_hd.docx": (
            "Hop dong {{SO_HOP_DONG}} ngay {{NGAY_HOP_DONG}} "
            "- chung thu {{NGAY_CHUNG_THU}} - {{TEN_KHACH_HANG}} - {{PHI_THAM_DINH}}"
        ),
        "mau_pyc.docx": "Phieu yeu cau {{TAI_SAN_THAM_DINH}}",
        "mau_bbnt.docx": "Bien ban nghiem thu {{SO_BIEN_BAN_NGHIEM_THU}}",
    }
    for name, text in templates.items():
        _write_docx_template(directory / name, text)


def _write_organization_templates(directory: Path) -> None:
    templates = {
        "hop_dong_vcb.docx": (
            "Hop dong to chuc {{SO_HOP_DONG}} - {{TEN_KHACH_HANG}} - {{MA_SO_THUE}} "
            "{{DIA_CHI_KHACH_HANG}} {{DIEN_THOAI_KHACH_HANG}} {{NGUOI_DAI_DIEN}} "
            "{{CHUC_VU_NGUOI_DAI_DIEN}} {{TAI_SAN_THAM_DINH}} {{MUC_DICH_THAM_DINH_DAY_DU}} "
            "{{SO_HOP_DONG_VAN_BAN}} {{PHI_THAM_DINH}} {{NGAY_HOP_DONG_NGAY}} "
            "{{NGAY_HOP_DONG_THANG}} {{NGAY_HOP_DONG_NAM}}"
        ),
        "hop_dong_vcb_tam_ung.docx": (
            "Hop dong tam ung {{TEN_KHACH_HANG}} {{DIA_CHI_KHACH_HANG}} {{DIEN_THOAI_KHACH_HANG}} "
            "{{MA_SO_THUE}} {{NGUOI_DAI_DIEN}} {{CHUC_VU_NGUOI_DAI_DIEN}} {{TAI_SAN_THAM_DINH}} "
            "{{MUC_DICH_THAM_DINH_DAY_DU}} {{SO_HOP_DONG_VAN_BAN}} {{PHI_THAM_DINH}} "
            "{{PHI_THAM_DINH_BANG_CHU}} {{TAM_UNG}} - {{TAM_UNG_BANG_CHU}} - "
            "{{CON_LAI_THANH_TOAN}} - {{CON_LAI_THANH_TOAN_BANG_CHU}} "
            "{{NGAY_HOP_DONG_NGAY}} {{NGAY_HOP_DONG_THANG}} {{NGAY_HOP_DONG_NAM}}"
        ),
        "bbtl_cong_ty.docx": (
            "Bien ban to chuc {{TEN_KHACH_HANG}} {{DIA_CHI_KHACH_HANG}} {{MA_SO_THUE}} "
            "{{NGUOI_DAI_DIEN}} - {{CHUC_VU_NGUOI_DAI_DIEN}} {{SO_BIEN_BAN_NGHIEM_THU}} "
            "{{SO_HOP_DONG_VAN_BAN}} {{NGAY_HOP_DONG}} {{PHI_THAM_DINH}} {{TAM_UNG}} "
            "{{CON_LAI_THANH_TOAN}} {{CON_LAI_THANH_TOAN_BANG_CHU}} {{NGAY}} {{THANG}} {{NAM}}"
        ),
        "de_nghi_thanh_toan.docx": (
            "De nghi thanh toan {{TEN_KHACH_HANG}} {{DIA_CHI_KHACH_HANG}} "
            "{{SO_DE_NGHI_THANH_TOAN}} {{SO_HOP_DONG_VAN_BAN}} {{PHI_THAM_DINH}} "
            "{{TAM_UNG}} {{CON_LAI_THANH_TOAN}} {{NGAY}} {{THANG}} {{NAM}}"
        ),
        "thu_chao_phi.docx": (
            "Thu chao phi {{TEN_KHACH_HANG}} {{TAI_SAN_THAM_DINH}} {{PHI_THAM_DINH}} "
            "{{PHI_THAM_DINH_BANG_CHU}} {{MUC_DICH_THAM_DINH_DAY_DU}} {{PHUONG_THUC_THANH_TOAN}} "
            "{{NGAY_HOP_DONG_NGAY}} {{NGAY_HOP_DONG_THANG}} {{NGAY_HOP_DONG_NAM}}"
        ),
    }
    for name, text in templates.items():
        _write_docx_template(directory / name, text)


def _write_required_templates_with_missing_placeholder(directory: Path, customer_type: str) -> dict[str, str]:
    missing_by_template: dict[str, str] = {}
    for template_name, required in TEMPLATE_REQUIREMENTS[customer_type].items():
        missing = sorted(required)[0]
        missing_by_template[template_name] = missing
        placeholders = [f"{{{{{name}}}}}" for name in sorted(required) if name != missing]
        _write_docx_template(directory / template_name, " ".join(placeholders))
    return missing_by_template


def _fake_pdf_export(word_paths: list[Path], *, soffice_path: str | Path) -> list[Path]:
    pdf_paths: list[Path] = []
    for word_path in word_paths:
        pdf_path = word_path.with_suffix(".pdf")
        pdf_path.write_bytes(b"%PDF-1.4\nfake pdf")
        pdf_paths.append(pdf_path)
    return pdf_paths


class CaseFiltersTests(unittest.TestCase):
    def test_filter_value_maps_all_to_empty_string(self) -> None:
        self.assertEqual(filter_value("Tất cả"), "")
        self.assertEqual(filter_value("03/2026"), "03/2026")

    def test_build_filters_normalizes_all_fields(self) -> None:
        filters = build_filters(
            selected_execution_month="03/2026",
            selected_payment_status="Tất cả",
            selected_case_status="Đang xử lý",
            selected_source="VCB",
            selected_customer_type="organization",
            selected_business_staff="Tất cả",
        )
        self.assertEqual(
            filters,
            {
                "execution_month": "03/2026",
                "payment_status": "",
                "case_status": "Đang xử lý",
                "source": "VCB",
                "customer_type": "organization",
                "business_staff": "",
            },
        )

    def test_chart_rows_and_data_are_stable(self) -> None:
        summary = {
            "target_month": "03/2026",
            "projected_current_month": 300,
            "paid_current_month": 200,
            "unpaid_current_month": 100,
            "previous_months": [
                {"Tháng": "02/2026", "Số hồ sơ": 2, "Doanh thu dự kiến": 200, "Đã thanh toán": 150, "Chưa thanh toán": 50},
                {"Tháng": "01/2026", "Số hồ sơ": 1, "Doanh thu dự kiến": 100, "Đã thanh toán": 100, "Chưa thanh toán": 0},
            ],
        }
        rows = build_chart_rows(summary, total_matches=3)
        self.assertEqual([row["Tháng"] for row in rows], ["01/2026", "02/2026", "03/2026"])

        chart_data = build_chart_data(rows)
        self.assertEqual(len(chart_data), 9)
        self.assertEqual(chart_data[0], {"Tháng": "01/2026", "Chỉ tiêu": "Doanh thu dự kiến", "Giá trị": 100})

    def test_chart_rows_use_revenue_case_count_when_available(self) -> None:
        summary = {
            "target_month": "03/2026",
            "case_count_current_month": 2,
            "projected_current_month": 300,
            "paid_current_month": 200,
            "unpaid_current_month": 100,
            "previous_months": [],
        }

        rows = build_chart_rows(summary, total_matches=3)

        self.assertEqual(rows[0]["Số hồ sơ"], 2)

    def test_export_scope_label_uses_readable_customer_type(self) -> None:
        label = export_scope_label(
            "ABC",
            {
                "execution_month": "03/2026",
                "payment_status": "Chưa thanh toán",
                "case_status": "",
                "source": "VCB",
                "customer_type": "organization",
                "business_staff": "Nguyễn A",
            },
        )
        self.assertEqual(label, "ABC | 03/2026 | Chưa thanh toán | VCB | Tổ chức | Nguyễn A")

    def test_get_unpaid_report_excludes_cancelled_cases_with_temp_sqlite(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            db_path = Path(tmpdir) / "cases.db"
            init_db(db_path)
            create_case(
                db_path,
                {
                    "contract_number": "HD-001",
                    "customer_info": "Khach A",
                    "customer_type": "individual",
                    "source": "VCB",
                    "business_staff": "NV A",
                    "execution_month": "03/2026",
                    "payment_status": "Chưa thanh toán",
                    "case_status": DEFAULT_CASE_STATUS,
                    "valuation_fee_number": "1000000",
                },
            )
            create_case(
                db_path,
                {
                    "contract_number": "HD-002",
                    "customer_info": "Khách B",
                    "execution_month": "03/2026",
                    "payment_status": "Chưa thanh toán",
                    "case_status": CANCELED_CASE_STATUS,
                    "valuation_fee_number": "2000000",
                },
            )

            report = get_unpaid_report(
                db_path,
                "",
                {
                    "execution_month": "03/2026",
                    "payment_status": "",
                    "case_status": "",
                    "source": "",
                    "customer_type": "",
                    "business_staff": "",
                },
            )

        self.assertEqual(report["count"], 1)
        self.assertEqual(report["total"], 1000000)
        self.assertIsInstance(report["rows"][0]["_id"], int)
        self.assertEqual(report["rows"][0]["Số hợp đồng"], "HD-001")

    def test_search_page_is_case_insensitive_for_vietnamese_text(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            db_path = Path(tmpdir) / "cases.db"
            init_db(db_path)
            create_case(
                db_path,
                {
                    "contract_number": "HD-DAKLAK-001",
                    "customer_info": "CÔNG TY ĐẮK LẮK",
                    "source": "VCB ĐẮK LẮK",
                    "payment_status": "Đã thanh toán",
                    "case_status": DEFAULT_CASE_STATUS,
                },
            )

            rows = search_page(
                db_path,
                "công ty đắk lắk",
                {
                    "execution_month": "",
                    "payment_status": "",
                    "case_status": "",
                    "source": "",
                    "customer_type": "",
                    "business_staff": "",
                },
                sort_field="contract_number",
                sort_direction="asc",
                page_size=10,
                page_number=1,
            )

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["customer_info"], "CÔNG TY ĐẮK LẮK")

    def test_search_page_filters_by_personal_note_only(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            db_path = Path(tmpdir) / "cases.db"
            init_db(db_path)
            create_case(
                db_path,
                {
                    "contract_number": "HD-001",
                    "customer_info": "Khach A",
                    "personal_note": "Cần kiểm tra lại pháp lý",
                    "payment_status": "Đã thanh toán",
                    "case_status": DEFAULT_CASE_STATUS,
                },
            )
            create_case(
                db_path,
                {
                    "contract_number": "HD-002",
                    "customer_info": "Khach B",
                    "personal_note": "Đã đủ hồ sơ",
                    "payment_status": "Đã thanh toán",
                    "case_status": DEFAULT_CASE_STATUS,
                },
            )

            rows = search_page(
                db_path,
                "",
                {
                    "execution_month": "",
                    "payment_status": "",
                    "case_status": "",
                    "source": "",
                    "customer_type": "",
                    "business_staff": "",
                },
                note_query="PHÁP LÝ",
                sort_field="contract_number",
                sort_direction="asc",
                page_size=10,
                page_number=1,
            )

        self.assertEqual([row["contract_number"] for row in rows], ["HD-001"])

    def test_execution_month_filter_options_show_newest_month_first(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            db_path = Path(tmpdir) / "cases.db"
            init_db(db_path)
            for month in ["03/2026", "04/2026", "12/2025"]:
                create_case(
                    db_path,
                    {
                        "contract_number": f"HD-{month}",
                        "customer_info": f"Khach {month}",
                        "execution_month": month,
                        "payment_status": "Đã thanh toán",
                        "case_status": DEFAULT_CASE_STATUS,
                    },
                )

            months = distinct_case_values(db_path, "execution_month")

        self.assertEqual(months[:3], ["04/2026", "03/2026", "12/2025"])

    def test_export_rows_for_filters_uses_current_filter_and_visible_columns(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            db_path = Path(tmpdir) / "cases.db"
            init_db(db_path)
            create_case(
                db_path,
                {
                    "contract_number": "HD-001",
                    "customer_info": "Khach A",
                    "execution_month": "03/2026",
                    "payment_status": "Đã thanh toán",
                    "case_status": DEFAULT_CASE_STATUS,
                    "valuation_fee_number": "3000000",
                },
            )
            create_case(
                db_path,
                {
                    "contract_number": "HD-002",
                    "customer_info": "Khach B",
                    "execution_month": "04/2026",
                    "payment_status": "Đã thanh toán",
                    "case_status": DEFAULT_CASE_STATUS,
                    "valuation_fee_number": "4000000",
                },
            )

            rows = export_rows_for_filters(
                db_path,
                "",
                {
                    "execution_month": "03/2026",
                    "payment_status": "",
                    "case_status": "",
                    "source": "",
                    "customer_type": "",
                    "business_staff": "",
                },
                visible_columns=["Số hợp đồng", "Thông tin khách hàng", "Phí thẩm định"],
                sort_field="contract_number",
                sort_direction="asc",
                export_count=10,
            )

        self.assertEqual(rows, [{"Số hợp đồng": "HD-001", "Thông tin khách hàng": "Khach A", "Phí thẩm định": "3.000.000"}])


class CaseExportsTests(unittest.TestCase):
    def test_case_folder_uses_contract_and_customer_without_id_prefix(self) -> None:
        folder = case_folder(
            Path("cases"),
            case_id=12,
            contract_number="010/2026/N04.1027/DN",
            customer_name="Ông Nguyễn Huy Hoàng - 0905222125 (Thu tiền)",
        )

        self.assertEqual(folder, Path("cases") / "010-2026-N04.1027-DN - Ông Nguyễn Huy Hoàng")

    def test_case_output_dir_preferences_round_trip(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            config_path = root / "config" / "case_output.json"
            default_dir = root / "default"
            selected_dir = root / "exports"

            self.assertEqual(load_case_output_dir(config_path, default_dir=default_dir), default_dir)
            saved = save_case_output_dir(config_path, selected_dir)

            self.assertEqual(saved, selected_dir)
            self.assertTrue(selected_dir.exists())
            self.assertEqual(load_case_output_dir(config_path, default_dir=default_dir), selected_dir)

    def test_document_action_error_validates_customer_type_and_pdf_tool(self) -> None:
        sample_case = {
            "customer_info": "Khách A",
            "customer_address": "Địa chỉ A",
            "asset_description": "Tài sản A",
            "valuation_purpose": "Mục đích A",
            "valuation_fee_number": 1000000,
        }
        self.assertEqual(
            document_action_error(
                case={"id": 1},
                expected_customer_type="individual",
                actual_customer_type="organization",
                template_errors=[],
            ),
            "Hồ sơ này đang là khách hàng tổ chức.",
        )
        self.assertEqual(
            document_action_error(
                case=sample_case,
                expected_customer_type="individual",
                actual_customer_type="individual",
                template_errors=[],
                require_pdf=True,
                soffice_path=None,
            ),
            "Không tìm thấy soffice.exe để xuất PDF.",
        )
        self.assertIsNone(
            document_action_error(
                case=sample_case,
                expected_customer_type="individual",
                actual_customer_type="individual",
                template_errors=[],
            )
        )

    def test_export_case_documents_routes_by_customer_type(self) -> None:
        case = {"id": 1}
        templates_dir = Path("templates")
        with patch("src.case_exports.export_individual_document_set", return_value=[Path("a.docx")]) as individual_export:
            self.assertEqual(export_case_documents(case, customer_type="individual", templates_dir=templates_dir), [Path("a.docx")])
            individual_export.assert_called_once()

        with patch("src.case_exports.export_organization_document_set", return_value=[Path("b.docx")]) as organization_export:
            self.assertEqual(export_case_documents(case, customer_type="organization", templates_dir=templates_dir), [Path("b.docx")])
            organization_export.assert_called_once()

    def test_export_case_documents_accepts_custom_output_directory(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            templates_dir = root / "templates"
            output_dir = root / "selected_output"
            _write_individual_templates(templates_dir)
            case = {
                "id": 7,
                "customer_type": "individual",
                "contract_number": "HD-CUSTOM-001",
                "customer_info": "Nguyen Van A",
                "asset_description": "Thua dat so 123.",
                "valuation_fee_number": "1500000",
            }

            paths = export_case_documents(
                case,
                customer_type="individual",
                templates_dir=templates_dir,
                case_files_dir=output_dir,
            )
            path_count = len(paths)
            paths_in_output_dir = all(str(path).startswith(str(output_dir)) for path in paths)
            paths_exist = all(path.exists() for path in paths)

        self.assertEqual(path_count, 3)
        self.assertTrue(paths_in_output_dir)
        self.assertTrue(paths_exist)

    def test_approve_pdf_exports_word_then_pdf(self) -> None:
        with (
            patch("src.case_exports.export_case_documents", return_value=[Path("a.docx")]) as export_word,
            patch("src.case_exports.export_docx_set_to_pdf", return_value=[Path("a.pdf")]) as export_pdf,
        ):
            word_paths, pdf_paths = approve_case_documents_pdf(
                {"id": 1},
                customer_type="individual",
                templates_dir=Path("templates"),
                soffice_path=Path("soffice.exe"),
            )
        self.assertEqual(word_paths, [Path("a.docx")])
        self.assertEqual(pdf_paths, [Path("a.pdf")])
        export_word.assert_called_once()
        export_pdf.assert_called_once_with([Path("a.docx")], soffice_path=Path("soffice.exe"))

    def test_package_case_documents_requires_folder(self) -> None:
        with self.assertRaises(ValueError):
            package_case_documents(None)

    def test_export_word_pdf_zip_from_sqlite_case(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            db_path = root / "cases.db"
            templates_dir = root / "templates"
            files_dir = root / "case_files"
            init_db(db_path)
            _write_individual_templates(templates_dir)

            case_id = create_case(
                db_path,
                {
                    "customer_type": "individual",
                    "contract_number": "HD-TEST-001",
                    "customer_info": "Nguyen Van A",
                    "customer_address": "12 Le Loi",
                    "asset_description": "Thua dat so 123.",
                    "valuation_purpose": "Vay von",
                    "source": "VCB",
                    "valuation_fee_number": "1500000",
                },
            )
            folder = case_folder(files_dir, case_id=case_id, contract_number="HD-TEST-001")
            original = root / "gcn.pdf"
            original.write_bytes(b"%PDF-1.4\noriginal")
            saved_original = save_original_file(original, "GCN.pdf", folder)
            update_case(
                db_path,
                case_id,
                {
                    "case_folder": str(folder),
                    "original_file_path": str(saved_original),
                },
            )
            case = get_case(db_path, case_id)
            assert case is not None

            word_paths = export_case_documents(case, customer_type="individual", templates_dir=templates_dir)
            with patch("src.case_exports.export_docx_set_to_pdf", side_effect=_fake_pdf_export):
                approved_word_paths, pdf_paths = approve_case_documents_pdf(
                    case,
                    customer_type="individual",
                    templates_dir=templates_dir,
                    soffice_path=Path("soffice.exe"),
                )
            zip_path = package_case_documents(folder)

            with ZipFile(zip_path) as archive:
                names = set(archive.namelist())
            word_count = len(word_paths)
            approved_word_count = len(approved_word_paths)
            pdf_count = len(pdf_paths)
            generated_files_exist = all(path.exists() for path in word_paths + pdf_paths)
            zip_exists = zip_path.exists()

        self.assertEqual(word_count, 3)
        self.assertEqual(approved_word_count, 3)
        self.assertEqual(pdf_count, 3)
        self.assertTrue(generated_files_exist)
        self.assertTrue(zip_exists)
        self.assertIn("originals/GCN.pdf", names)
        self.assertEqual(len([name for name in names if "/" not in name and name.endswith(".docx")]), 3)
        self.assertEqual(len([name for name in names if "/" not in name and name.endswith(".pdf")]), 3)
        self.assertFalse(any(name.startswith("documents/") for name in names))

    def test_exported_word_renders_placeholders_with_case_values(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            db_path = root / "cases.db"
            templates_dir = root / "templates"
            files_dir = root / "case_files"
            init_db(db_path)
            _write_individual_templates(templates_dir)
            case_id = create_case(
                db_path,
                {
                    "customer_type": "individual",
                    "contract_number": "HD-PLACE-001",
                    "contract_date": "2026-04-25",
                    "certificate_date": "26/04/2026",
                    "customer_info": "Nguyen Van A",
                    "asset_description": "Thua dat so 123.",
                    "valuation_fee_number": "1500000",
                },
            )
            folder = case_folder(files_dir, case_id=case_id, contract_number="HD-PLACE-001")
            update_case(db_path, case_id, {"case_folder": str(folder)})
            case = get_case(db_path, case_id)
            assert case is not None

            word_paths = export_case_documents(case, customer_type="individual", templates_dir=templates_dir)
            hop_dong_path = next(path for path in word_paths if path.name.endswith("_hop_dong.docx"))
            content = read_docx_text(hop_dong_path)

        self.assertIn("HD-PLACE-001", content)
        self.assertIn("25/04/2026", content)
        self.assertIn("26/04/2026", content)
        self.assertIn("Nguyen Van A", content)
        self.assertIn("1.500.000", content)
        self.assertNotIn("{{", content)
        self.assertNotIn("}}", content)

    def test_docx_preview_html_keeps_basic_word_formatting(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "template.docx"
            document = Document()
            paragraph = document.add_paragraph()
            bold = paragraph.add_run("Khách hàng: ")
            bold.bold = True
            italic = paragraph.add_run("{{TEN_KHACH_HANG}}")
            italic.italic = True
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Số HĐ"
            table.cell(0, 1).text = "{{SO_HOP_DONG}}"
            document.save(path)

            rendered = render_docx_preview_html(
                path,
                {
                    "customer_info": "Nguyen Van A",
                    "contract_number": "HD-001",
                },
            )

        self.assertIn("font-weight:700", rendered)
        self.assertIn("font-style:italic", rendered)
        self.assertIn("<table>", rendered)
        self.assertIn("Nguyen Van A", rendered)
        self.assertIn("HD-001", rendered)
        self.assertNotIn("{{", rendered)

    def test_placeholder_context_does_not_append_internal_source_or_duplicate_bank(self) -> None:
        context = build_placeholder_context(
            {
                "valuation_purpose": "Làm cơ sở tham khảo để thế chấp vay vốn tại VP Bank",
                "source": "KHN",
            }
        )
        self.assertEqual(
            context["MUC_DICH_THAM_DINH_DAY_DU"],
            "Làm cơ sở tham khảo để thế chấp vay vốn tại VP\xa0Bank.",
        )

        context = build_placeholder_context(
            {
                "valuation_purpose": "Làm cơ sở tham khảo để thế chấp vay vốn tại Sacombank",
                "source": "Sacombank Kon Tum - Mr. Luân - 0967097789",
            }
        )
        self.assertEqual(
            context["MUC_DICH_THAM_DINH_DAY_DU"],
            "Làm cơ sở tham khảo để thế chấp vay vốn tại Sacombank.",
        )

    def test_placeholder_context_for_request_form_uses_land_address_short_purpose_and_contract_date(self) -> None:
        context = build_placeholder_context(
            {
                "asset_description": (
                    "Giá trị quyền sử dụng đất tại Thửa đất số 136, tờ bản đồ số 325, "
                    "tại địa chỉ thôn Vinh Bình, xã Cam Lâm, tỉnh Khánh Hòa."
                ),
                "customer_address": "44 Nguyễn Du",
                "valuation_purpose": "Làm cơ sở tham khảo để thế chấp vay vốn tại VP Bank",
                "source": "KHN",
                "contract_date": "25/04/2026",
            }
        )

        self.assertEqual(context["DIA_CHI_TAI_SAN"], "thôn Vinh Bình, xã Cam Lâm, tỉnh Khánh Hòa")
        self.assertEqual(context["MUC_DICH_THAM_DINH_RUT_GON"], "thế chấp vay vốn tại VP\xa0Bank.")
        self.assertEqual(context["NGAY_HOP_DONG_PLEIKU"], "Pleiku, Ngày 25 tháng 04 năm 2026")

    def test_request_form_short_purpose_does_not_include_source_contact(self) -> None:
        context = build_placeholder_context(
            {
                "valuation_purpose": "Làm cơ sở tham khảo để xử lý tài sản đảm bảo tại ngân hàng",
                "source": "BIDV Phố Núi - Ms Quỳnh Anh - 0975172634",
            }
        )

        self.assertEqual(
            context["MUC_DICH_THAM_DINH_RUT_GON"],
            "xử lý tài sản đảm bảo tại ngân hàng.",
        )
        self.assertEqual(
            context["MUC_DICH_THAM_DINH_DAY_DU"],
            "Làm cơ sở tham khảo để xử lý tài sản đảm bảo tại ngân hàng.",
        )
        self.assertNotIn("BIDV", context["MUC_DICH_THAM_DINH_RUT_GON"])
        self.assertNotIn("BIDV", context["MUC_DICH_THAM_DINH_DAY_DU"])
        self.assertNotIn("0975172634", context["MUC_DICH_THAM_DINH_RUT_GON"])
        self.assertNotIn("0975172634", context["MUC_DICH_THAM_DINH_DAY_DU"])

    def test_placeholder_context_ignores_excel_formula_error_fee_words(self) -> None:
        context = build_placeholder_context(
            {
                "valuation_fee_number": "4400000",
                "valuation_fee_words": "#NAME?",
            }
        )

        self.assertEqual(context["PHI_THAM_DINH"], "4.400.000")
        self.assertNotEqual(context["PHI_THAM_DINH_BANG_CHU"], "#NAME?")
        self.assertEqual(context["PHI_THAM_DINH_BANG_CHU"], "Bốn triệu bốn trăm ngàn đồng chẵn")

    def test_sample_templates_do_not_repeat_phone_or_leave_asset_dot_leaders(self) -> None:
        sample_root = Path("samples/templates")
        violations: list[str] = []
        for path in sample_root.rglob("*.docx"):
            text = read_docx_text(path)
            if "{{DIEN_THOAI_KHACH_HANG}} {{DIEN_THOAI_KHACH_HANG}}" in text:
                violations.append(f"{path}: lặp placeholder điện thoại")
            if "{{TAI_SAN_THAM_DINH}}..." in text:
                violations.append(f"{path}: còn dấu chấm đệm sau tài sản")

        self.assertEqual(violations, [])

    def test_sample_request_form_uses_short_purpose_and_contract_date_placeholders(self) -> None:
        text = read_docx_text(Path("samples/templates/individual/mau_pyc.docx"))

        self.assertIn("Địa điểm khảo sát tài sản: {{DIA_CHI_TAI_SAN}}.", text)
        self.assertNotIn("{{TAI_SAN_THAM_DINH}}.", text)
        self.assertIn("mục đích {{MUC_DICH_THAM_DINH_RUT_GON}}", text)
        self.assertIn("{{NGAY_HOP_DONG_PLEIKU}}", text)
        self.assertNotIn("{{NGAY_LAP_PLEIKU}}", text)

    def test_sample_organization_acceptance_uses_contract_date_placeholder(self) -> None:
        text = read_docx_text(Path("samples/templates/organization/bbtl_cong_ty.docx"))

        self.assertIn("{{NGAY_HOP_DONG}}", text)
        self.assertNotIn("06/10/2025", text)
        self.assertIn("NGAY_HOP_DONG", TEMPLATE_REQUIREMENTS["organization"]["bbtl_cong_ty.docx"])

    def test_render_docx_template_preserves_run_formatting_when_replacing_placeholders(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            template = root / "template.docx"
            output = root / "output.docx"
            document = Document()
            paragraph = document.add_paragraph()
            label = paragraph.add_run("Khách hàng: ")
            label.bold = True
            value = paragraph.add_run("{{TEN_KHACH_HANG}}")
            value.italic = True
            document.save(template)

            render_docx_template(
                template,
                output,
                {
                    "id": 1,
                    "customer_info": "Nguyen Van A",
                },
            )
            rendered = Document(output)
            runs = rendered.paragraphs[0].runs

        self.assertEqual(runs[0].text, "Khách hàng: ")
        self.assertTrue(runs[0].bold)
        self.assertEqual(runs[1].text, "Nguyen Van A")
        self.assertTrue(runs[1].italic)

    def test_render_docx_template_handles_placeholder_split_across_runs(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            template = root / "template.docx"
            output = root / "output.docx"
            document = Document()
            paragraph = document.add_paragraph()
            first = paragraph.add_run("{{TEN")
            first.italic = True
            second = paragraph.add_run("_KHACH_HANG}}")
            second.bold = True
            document.save(template)

            render_docx_template(
                template,
                output,
                {
                    "id": 1,
                    "customer_info": "Nguyen Van A",
                },
            )
            rendered = Document(output)
            runs = rendered.paragraphs[0].runs

        self.assertEqual("".join(run.text for run in runs), "Nguyen Van A")
        self.assertTrue(runs[0].italic)

    def test_preview_and_exported_word_match_after_render(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            db_path = root / "cases.db"
            templates_dir = root / "templates"
            files_dir = root / "case_files"
            init_db(db_path)
            _write_individual_templates(templates_dir)
            case_id = create_case(
                db_path,
                {
                    "customer_type": "individual",
                    "contract_number": "HD-COMPARE-001",
                    "customer_info": "Nguyen Van B",
                    "customer_address": "34 Tran Phu",
                    "asset_description": "Can ho chung cu.",
                    "valuation_purpose": "Vay von",
                    "source": "VCB",
                    "valuation_fee_number": "2200000",
                },
            )
            folder = case_folder(files_dir, case_id=case_id, contract_number="HD-COMPARE-001")
            update_case(db_path, case_id, {"case_folder": str(folder)})
            case = get_case(db_path, case_id)
            assert case is not None

            export_case_documents(case, customer_type="individual", templates_dir=templates_dir)
            comparisons = compare_case_documents(case, customer_type="individual", templates_dir=templates_dir)

        self.assertEqual(len(comparisons), 3)
        self.assertTrue(all(item["matched"] for item in comparisons))
        self.assertTrue(all(item["reason"] == "" for item in comparisons))
        for item in comparisons:
            self.assertEqual(item["preview_content"], item["exported_content"])

    def test_export_organization_word_pdf_zip_from_sqlite_case(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            db_path = root / "cases.db"
            templates_dir = root / "organization_templates"
            files_dir = root / "case_files"
            init_db(db_path)
            _write_organization_templates(templates_dir)

            case_id = create_case(
                db_path,
                {
                    "customer_type": "organization",
                    "contract_number": "HD-ORG-001",
                    "customer_info": "Cong ty TNHH ABC",
                    "customer_address": "99 Nguyen Tat Thanh",
                    "asset_description": "Nha xuong.",
                    "valuation_purpose": "Vay von",
                    "source": "VCB",
                    "valuation_fee_number": "2500000",
                    "advance_payment": "500000",
                    "tax_code": "6001234567",
                    "representative_name": "Tran Van B",
                    "representative_position": "Giam doc",
                },
            )
            folder = case_folder(files_dir, case_id=case_id, contract_number="HD-ORG-001")
            original = root / "org_gcn.pdf"
            original.write_bytes(b"%PDF-1.4\norganization original")
            saved_original = save_original_file(original, "GCN_To_Chuc.pdf", folder)
            update_case(
                db_path,
                case_id,
                {
                    "case_folder": str(folder),
                    "original_file_path": str(saved_original),
                },
            )
            case = get_case(db_path, case_id)
            assert case is not None

            word_paths = export_case_documents(case, customer_type="organization", templates_dir=templates_dir)
            with patch("src.case_exports.export_docx_set_to_pdf", side_effect=_fake_pdf_export):
                approved_word_paths, pdf_paths = approve_case_documents_pdf(
                    case,
                    customer_type="organization",
                    templates_dir=templates_dir,
                    soffice_path=Path("soffice.exe"),
                )
            zip_path = package_case_documents(folder)

            with ZipFile(zip_path) as archive:
                names = set(archive.namelist())
            word_count = len(word_paths)
            approved_word_count = len(approved_word_paths)
            pdf_count = len(pdf_paths)
            generated_files_exist = all(path.exists() for path in word_paths + pdf_paths)

        self.assertEqual(word_count, 4)
        self.assertEqual(approved_word_count, 4)
        self.assertEqual(pdf_count, 4)
        self.assertTrue(generated_files_exist)
        self.assertIn("originals/GCN_To_Chuc.pdf", names)
        self.assertEqual(len([name for name in names if "/" not in name and name.endswith(".docx")]), 4)
        self.assertEqual(len([name for name in names if "/" not in name and name.endswith(".pdf")]), 4)
        self.assertFalse(any(name.startswith("documents/") for name in names))

    def test_export_organization_contract_with_advance_template(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            templates_dir = root / "organization_templates"
            files_dir = root / "case_files"
            _write_organization_templates(templates_dir)
            case = {
                "id": 9,
                "customer_type": "organization",
                "organization_contract_payment_method": "advance",
                "contract_number": "HD-ADV-001",
                "contract_date": "25/05/2026",
                "customer_info": "Cong ty TNHH ABC",
                "customer_address": "99 Nguyen Tat Thanh",
                "asset_description": "Nha xuong.",
                "valuation_purpose": "Vay von",
                "source": "VCB",
                "valuation_fee_number": "2500000",
                "advance_payment": "500000",
                "tax_code": "6001234567",
                "representative_name": "Tran Van B",
                "representative_position": "Giam doc",
            }

            word_paths = export_case_documents(
                case,
                customer_type="organization",
                templates_dir=templates_dir,
                case_files_dir=files_dir,
            )
            hop_dong_path = next(path for path in word_paths if path.name.endswith("_hop_dong_to_chuc.docx"))
            content = read_docx_text(hop_dong_path)
            thu_chao_phi_path = next(path for path in word_paths if path.name.endswith("_thu_chao_phi.docx"))
            thu_chao_phi = read_docx_text(thu_chao_phi_path)

        self.assertIn("Hop dong tam ung Cong ty TNHH ABC", content)
        self.assertIn("500.000", content)
        self.assertIn("Năm trăm ngàn đồng chẵn", content)
        self.assertIn("2.000.000", content)
        self.assertIn("Hai triệu đồng chẵn", content)
        self.assertNotIn("{{", content)
        self.assertIn("Bên A thanh toán cho Bên B làm 02 đợt", thu_chao_phi)
        self.assertIn("25 05 2026", thu_chao_phi)
        self.assertIn("500.000", thu_chao_phi)
        self.assertIn("2.000.000", thu_chao_phi)
        self.assertNotIn("{{", thu_chao_phi)

    def test_export_vcb_gia_lai_customer_uses_dedicated_templates(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            templates_dir = root / "organization_templates"
            files_dir = root / "case_files"
            _write_organization_templates(templates_dir)
            special_dir = templates_dir / "vcb_gia_lai"
            special_templates = {
                "hop_dong_vcb.docx": "VCB Gia Lai contract {{SO_HOP_DONG}} {{TEN_KHACH_HANG}}",
                "bbtl_cong_ty.docx": "VCB Gia Lai acceptance {{SO_HOP_DONG}}",
                "de_nghi_thanh_toan.docx": "VCB Gia Lai payment {{CON_LAI_THANH_TOAN}}",
                "thu_chao_phi.docx": "VCB Gia Lai quote {{THANG_NAM_HOP_DONG}}",
            }
            for name, text in special_templates.items():
                _write_docx_template(special_dir / name, text)
            case = {
                "id": 10,
                "customer_type": "organization",
                "contract_number": "010/2026/N02-0522/DN",
                "contract_date": "25/02/2026",
                "customer_info": (
                    "NGÂN HÀNG THƯƠNG MẠI CỔ PHẦN NGOẠI THƯƠNG VIỆT NAM - "
                    "CHI NHÁNH GIA LAI"
                ),
                "valuation_fee_number": "25000000",
                "advance_payment": "0",
            }

            word_paths = export_case_documents(
                case,
                customer_type="organization",
                templates_dir=templates_dir,
                case_files_dir=files_dir,
            )
            contract_path = next(path for path in word_paths if path.name.endswith("_hop_dong_to_chuc.docx"))
            quote_path = next(path for path in word_paths if path.name.endswith("_thu_chao_phi.docx"))
            contract_text = read_docx_text(contract_path)
            quote_text = read_docx_text(quote_path)

        self.assertIn("VCB Gia Lai contract", contract_text)
        self.assertIn("Tháng 02 năm 2026", quote_text)

    def test_vcb_gia_lai_fixed_template_does_not_require_organization_contact_fields(self) -> None:
        case = {
            "customer_type": "organization",
            "customer_info": (
                "NGÂN HÀNG THƯƠNG MẠI CỔ PHẦN NGOẠI THƯƠNG VIỆT NAM – "
                "CHI NHÁNH GIA LAI (Thu tiền)"
            ),
            "customer_address": "5.02 Lô G KDC Miếu Nổi",
            "asset_description": "Tài sản thẩm định.",
            "valuation_purpose": "Làm cơ sở tham khảo",
            "valuation_fee_number": "3300000",
            "tax_code": "",
            "representative_name": "",
            "representative_position": "",
        }

        self.assertEqual(missing_mandatory_data_fields(case, "organization"), [])

    def test_collect_template_errors_reports_missing_required_individual_templates(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            templates_dir = Path(tmpdir)
            _write_docx_template(templates_dir / "mau_hd.docx", "Hop dong {{TEN_KHACH_HANG}}")

            errors = collect_template_errors(templates_dir, "individual")

        joined = "\n".join(errors)
        self.assertIn("mau_bbnt.docx", joined)
        self.assertIn("mau_pyc.docx", joined)
        self.assertIn("file template", joined)

    def test_collect_template_errors_reports_missing_required_organization_templates(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            templates_dir = Path(tmpdir)
            _write_docx_template(templates_dir / "hop_dong_vcb.docx", "Hop dong {{TEN_KHACH_HANG}}")

            errors = collect_template_errors(templates_dir, "organization")

        joined = "\n".join(errors)
        self.assertIn("bbtl_cong_ty.docx", joined)
        self.assertIn("de_nghi_thanh_toan.docx", joined)
        self.assertIn("thu_chao_phi.docx", joined)
        self.assertIn("file template", joined)

    def test_collect_template_errors_reports_all_required_files_when_folder_is_empty(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            templates_dir = Path(tmpdir)

            individual_errors = "\n".join(collect_template_errors(templates_dir, "individual"))
            organization_errors = "\n".join(collect_template_errors(templates_dir, "organization"))

        self.assertIn("mau_hd.docx", individual_errors)
        self.assertIn("mau_pyc.docx", individual_errors)
        self.assertIn("mau_bbnt.docx", individual_errors)
        self.assertIn("hop_dong_vcb.docx", organization_errors)
        self.assertIn("bbtl_cong_ty.docx", organization_errors)
        self.assertIn("de_nghi_thanh_toan.docx", organization_errors)
        self.assertIn("thu_chao_phi.docx", organization_errors)

    def test_collect_template_errors_reports_missing_placeholders_in_each_individual_template(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            templates_dir = Path(tmpdir)
            missing_by_template = _write_required_templates_with_missing_placeholder(templates_dir, "individual")

            errors = collect_template_errors(templates_dir, "individual")

        joined = "\n".join(errors)
        self.assertNotIn("file template", joined)
        for template_name, missing in missing_by_template.items():
            self.assertIn(template_name, joined)
            self.assertIn(f"{{{{{missing}}}}}", joined)

    def test_collect_template_errors_reports_missing_placeholders_in_each_organization_template(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            templates_dir = Path(tmpdir)
            missing_by_template = _write_required_templates_with_missing_placeholder(templates_dir, "organization")

            errors = collect_template_errors(templates_dir, "organization")

        joined = "\n".join(errors)
        self.assertNotIn("file template", joined)
        for template_name, missing in missing_by_template.items():
            self.assertIn(template_name, joined)
            self.assertIn(f"{{{{{missing}}}}}", joined)

    def test_package_case_documents_skips_existing_zip_files_in_package_folder(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            folder = root / "case"
            package_dir = folder / "package"
            folder.mkdir(parents=True)
            package_dir.mkdir(parents=True)
            (folder / "hop_dong.docx").write_bytes(b"docx")
            old_zip = package_dir / "old_package.zip"
            old_zip.write_bytes(b"old zip should not be nested")

            zip_path = package_case_documents(folder)

            with ZipFile(zip_path) as archive:
                names = set(archive.namelist())

        self.assertIn("hop_dong.docx", names)
        self.assertNotIn("package/old_package.zip", names)
        self.assertFalse(any(name.startswith("package/") for name in names))


if __name__ == "__main__":
    unittest.main()
