from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from src.template_manager import (
    append_template_history,
    create_template_snapshot,
    list_editable_blocks,
    read_docx_text,
    read_template_history,
    restore_template_from_snapshot,
    update_template_blocks,
)


class TemplateHistoryTests(unittest.TestCase):
    def test_restore_template_from_snapshot_restores_docx_content(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            template_path = root / "mau_hd.docx"
            versions_dir = root / "versions"

            document = Document()
            document.add_paragraph("Hop dong {{TEN_KHACH_HANG}}")
            document.save(template_path)

            snapshot_path = create_template_snapshot(template_path, versions_dir, reason="before_edit")
            changes = update_template_blocks(template_path, {"body-p-1": "Da sua {{TEN_KHACH_HANG}}"})
            self.assertEqual(len(changes), 1)
            self.assertEqual(read_docx_text(template_path), "Da sua {{TEN_KHACH_HANG}}")

            restore_template_from_snapshot(template_path, snapshot_path)
            self.assertEqual(read_docx_text(template_path), "Hop dong {{TEN_KHACH_HANG}}")

    def test_append_and_read_template_history_filters_by_template(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            history_path = root / "history.jsonl"
            template_a = root / "a.docx"
            template_b = root / "b.docx"
            template_a.write_text("placeholder", encoding="utf-8")
            template_b.write_text("placeholder", encoding="utf-8")

            append_template_history(
                history_path,
                template_path=template_a,
                editor_name="Tester",
                action="edit_block",
                details={"block_id": "body-p-1"},
            )
            append_template_history(
                history_path,
                template_path=template_b,
                editor_name="Tester",
                action="lock_template",
                details={"status": "locked"},
            )

            rows = read_template_history(history_path, template_path=template_a)

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["template_name"], "a.docx")
        self.assertEqual(rows[0]["action"], "edit_block")
        self.assertEqual(rows[0]["details"], {"block_id": "body-p-1"})

    def test_update_template_blocks_updates_body_table_header_and_footer(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            template_path = Path(tmpdir) / "template.docx"
            document = Document()
            document.add_paragraph("Body {{TEN_KHACH_HANG}}")
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).paragraphs[0].text = "Table {{PHI_THAM_DINH}}"
            section = document.sections[0]
            section.header.paragraphs[0].text = "Header {{SO_HOP_DONG}}"
            section.footer.paragraphs[0].text = "Footer {{NGAY}}"
            document.save(template_path)

            blocks = {block["block_id"]: block["text"] for block in list_editable_blocks(template_path)}
            self.assertEqual(blocks["body-p-1"], "Body {{TEN_KHACH_HANG}}")
            self.assertEqual(blocks["table-1-r-1-c-1-p-1"], "Table {{PHI_THAM_DINH}}")
            self.assertEqual(blocks["header-1-p-1"], "Header {{SO_HOP_DONG}}")
            self.assertEqual(blocks["footer-1-p-1"], "Footer {{NGAY}}")

            changes = update_template_blocks(
                template_path,
                {
                    "body-p-1": "Body updated {{TEN_KHACH_HANG}}",
                    "table-1-r-1-c-1-p-1": "Table updated {{PHI_THAM_DINH}}",
                    "header-1-p-1": "Header updated {{SO_HOP_DONG}}",
                    "footer-1-p-1": "Footer updated {{NGAY}}",
                },
            )
            updated_blocks = {block["block_id"]: block["text"] for block in list_editable_blocks(template_path)}

        self.assertEqual(len(changes), 4)
        self.assertEqual(updated_blocks["body-p-1"], "Body updated {{TEN_KHACH_HANG}}")
        self.assertEqual(updated_blocks["table-1-r-1-c-1-p-1"], "Table updated {{PHI_THAM_DINH}}")
        self.assertEqual(updated_blocks["header-1-p-1"], "Header updated {{SO_HOP_DONG}}")
        self.assertEqual(updated_blocks["footer-1-p-1"], "Footer updated {{NGAY}}")


if __name__ == "__main__":
    unittest.main()
