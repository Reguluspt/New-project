import sys
try:
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')
except Exception:
    pass

import imaplib
import email
from email.header import decode_header
import smtplib
import html
import shutil
import tempfile
import base64
from datetime import datetime, timedelta
from email.message import EmailMessage
from email.utils import formataddr, parseaddr
import os
import re
import asyncio
from pathlib import Path
from .mail_service import load_gmail_smtp_settings, _parse_email_list, _clean_header

DEFAULT_PHATHANH_RECIPIENT = (
    "CÔNG TY CỔ PHẦN THẨM ĐỊNH GIÁ THẾ KỶ - VP TẠI GIA LAI\n"
    "Địa chỉ: 90/60/3 Trường Chinh, phường Pleiku, tỉnh Gia Lai\n"
    "Điện thoại 0905226968"
)

DEFAULT_PHATHANH_DOCX_TEMPLATE = r"D:\Nhap\Form Phat Hanh Chung Thu.docx"


def gemini_keys_with_backups(primary_key: str | None = None, backup_keys: str | None = None) -> list[str]:
    keys: list[str] = []
    backup_value = backup_keys if backup_keys is not None else os.getenv("GEMINI_BACKUP_KEYS", "")
    for key in [primary_key or os.getenv("GEMINI_API_KEY", ""), *backup_value.split(",")]:
        clean_key = str(key or "").strip()
        if clean_key and clean_key not in keys:
            keys.append(clean_key)
    return keys


def build_phathanh_browser_tools(html_body: str):
    from browser_use import Tools
    from browser_use.agent.views import ActionResult

    tools = Tools()

    @tools.action("Insert the prepared certificate release HTML into the currently open OWA reply editor.")
    async def insert_phathanh_html(browser_session):
        page = await browser_session.must_get_current_page()
        result = await page.evaluate(
                r"""
            async (html) => {
              const isVisible = (el) => {
                const rect = el.getBoundingClientRect();
                const style = window.getComputedStyle(el);
                return rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none';
              };

              const candidates = Array.from(document.querySelectorAll([
                '[contenteditable="true"]',
                '[role="textbox"]',
                '[aria-label*="Nội dung"]',
                '[aria-label*="Message body"]',
                '[aria-label*="body"]'
              ].join(','))).filter(isVisible);

              let editor = null;
              const active = document.activeElement;
              if (active && candidates.includes(active)) {
                editor = active;
              }
              if (!editor && active && active.closest) {
                editor = candidates.find((candidate) => candidate.contains(active) || active.contains(candidate));
              }
              if (!editor) {
                editor = candidates[candidates.length - 1];
              }
              if (!editor) {
                return { ok: false, reason: 'reply editor not found' };
              }

              editor.focus();
              const selection = window.getSelection();
              const range = document.createRange();
              range.selectNodeContents(editor);
              range.collapse(false);
              selection.removeAllRanges();
              selection.addRange(range);

              let inserted = false;
              try {
                const clipboardData = new DataTransfer();
                clipboardData.setData('text/html', html);
                clipboardData.setData('text/plain', html.replace(/<[^>]+>/g, ' '));
                inserted = editor.dispatchEvent(new ClipboardEvent('paste', {
                  bubbles: true,
                  cancelable: true,
                  clipboardData,
                }));
              } catch (error) {
                inserted = false;
              }
              if (inserted) {
                const renderedText = (editor.innerText || editor.textContent || '').trim();
                inserted = renderedText.includes('Dear all') || editor.innerHTML.includes('Dear all');
              }
              try {
                inserted = inserted || document.execCommand('insertHTML', false, html);
              } catch (error) {
              }
              if (!inserted) {
                editor.innerHTML = html;
              }

              editor.dispatchEvent(new InputEvent('input', { bubbles: true, inputType: 'insertHTML', data: html }));
              editor.dispatchEvent(new Event('change', { bubbles: true }));

              const renderedText = (editor.innerText || editor.textContent || '').trim();
              return {
                ok: renderedText.includes('Dear all') || editor.innerHTML.includes('Dear all'),
                reason: renderedText.slice(0, 160),
              };
            }
            """,
            html_body,
        )
        if result.get("ok"):
            return ActionResult(extracted_content="Inserted phathanh HTML into the OWA reply editor.")
        return ActionResult(error=f"Could not insert phathanh HTML: {result.get('reason')}")

    return tools


def format_phathanh_recipient_html(recipient: str | None) -> str:
    from src.email_utils import format_recipient_info

    recipient_clean = format_recipient_info(recipient or DEFAULT_PHATHANH_RECIPIENT)
    recipient_html = html.escape(recipient_clean).replace("\n", "<br>")
    if "Địa chỉ:" in recipient_html:
        recipient_html = recipient_html.replace("Địa chỉ:", "<strong>Địa chỉ:</strong>")
    elif "Địa chỉ" in recipient_html:
        recipient_html = recipient_html.replace("Địa chỉ", "<strong>Địa chỉ</strong>")

    if "Điện thoại:" in recipient_html:
        recipient_html = recipient_html.replace("Điện thoại:", "<strong>Điện thoại:</strong>")
    elif "Điện thoại" in recipient_html:
        recipient_html = recipient_html.replace("Điện thoại", "<strong>Điện thoại</strong>")
    return recipient_html


def phathanh_send_mode() -> str:
    return os.getenv("PHATHANH_SEND_MODE", "playwright").strip().casefold()


def phathanh_docx_template_path() -> Path:
    configured_path = Path(os.getenv("PHATHANH_DOCX_TEMPLATE", DEFAULT_PHATHANH_DOCX_TEMPLATE))
    if configured_path.exists():
        return configured_path

    project_data_path = Path(__file__).resolve().parent.parent / "data" / "form_phat_hanh_chung_thu.docx"
    if project_data_path.exists():
        return project_data_path

    return configured_path


def phathanh_logo_data_uri() -> str:
    logo_path = Path(__file__).resolve().parent / "templates" / "logo.jpg"
    if not logo_path.exists():
        return "cid:logo_cenvalue"
    encoded = base64.b64encode(logo_path.read_bytes()).decode("ascii")
    return f"data:image/jpeg;base64,{encoded}"


def _phathanh_case_values(case: dict, recipient: str | None) -> dict[str, str]:
    now = datetime.now()
    recipient_clean = recipient or DEFAULT_PHATHANH_RECIPIENT
    return {
        "customer_name": str(case.get("customer_info") or case.get("owner_name") or "N/A"),
        "customer_address": str(case.get("customer_address") or case.get("dia_chi_thua_dat") or "N/A"),
        "recipient_info": recipient_clean,
        "date_receive": (now + timedelta(days=1)).strftime("%d/%m/%Y"),
        "date_payment": now.strftime("%d/%m/%Y"),
        "personal_note": str(case.get("personal_note") or ""),
    }


def _set_docx_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"


def create_phathanh_docx_for_case(case: dict, recipient: str | None = None) -> Path:
    from docx import Document

    template_path = phathanh_docx_template_path()
    if not template_path.exists():
        raise FileNotFoundError(f"Thiếu form phát hành chứng thư: {template_path}")

    output_dir = Path(tempfile.gettempdir()) / "century_phathanh"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"phathanh_{case.get('id') or 'case'}_{datetime.now().strftime('%Y%m%d%H%M%S%f')}.docx"
    shutil.copy2(template_path, output_path)

    values = _phathanh_case_values(case, recipient)
    doc = Document(output_path)
    if not doc.tables:
        doc.save(output_path)
        return output_path

    table = doc.tables[0]
    _set_docx_cell_text(table.cell(2, 0), f"Khách hàng yêu cầu*: {values['customer_name']}")
    _set_docx_cell_text(table.cell(3, 0), f"Địa chỉ liên lạc*: {values['customer_address']}")
    _set_docx_cell_text(table.cell(6, 2), values["recipient_info"])
    _set_docx_cell_text(table.cell(6, 3), values["date_receive"])
    _set_docx_cell_text(table.cell(6, 6), values["date_payment"])
    if values["personal_note"]:
        _set_docx_cell_text(table.cell(8, 0), f"Ghi chú: {values['personal_note']}")
    doc.save(output_path)
    return output_path


def phathanh_docx_to_email_html(docx_path: Path) -> str:
    from docx import Document

    doc = Document(docx_path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(
                '<p style="margin:0 0 6px 0;font-family:Arial,sans-serif;font-size:13px;">'
                f"{html.escape(text)}</p>"
            )

    col_widths = [115, 30, 309, 122, 113, 235, 84]
    table_style = (
        "width:1008px;max-width:100%;border-collapse:collapse;margin:8px 0;"
        "font-family:Arial,sans-serif;font-size:12px;table-layout:fixed;color:#000;"
    )
    cell_style = (
        "border:1px solid #000;padding:3px 6px;vertical-align:middle;"
        "line-height:1.18;background:#fff;word-break:normal;"
    )
    for table in doc.tables:
        parts.append(f'<table style="{table_style}">')
        parts.append("<colgroup>")
        for width in col_widths:
            parts.append(f'<col style="width:{width}px">')
        parts.append("</colgroup>")
        for row_index, row in enumerate(table.rows):
            parts.append("<tr>")
            cell_index = 0
            cells = row.cells
            while cell_index < len(cells):
                cell = cells[cell_index]
                colspan = 1
                while (
                    cell_index + colspan < len(cells)
                    and cells[cell_index + colspan]._tc is cell._tc
                ):
                    colspan += 1
                cell_text = "<br>".join(
                    html.escape(line) for line in cell.text.splitlines()
                ) or "&nbsp;"
                style = cell_style
                if row_index in (0, 1, 4, 5):
                    style += "font-weight:bold;text-align:center;background:#E2EFDA;"
                if row_index == 0:
                    style += "background:#FFC000;"
                if row_index in (2, 3):
                    style += "text-align:left;"
                if row_index in (6, 7):
                    if cell_index in (1, 3, 6):
                        style += "text-align:center;"
                    else:
                        style += "text-align:left;"
                if row_index == 8:
                    style += "font-weight:bold;"
                colspan_attr = f' colspan="{colspan}"' if colspan > 1 else ""
                parts.append(f'<td{colspan_attr} style="{style}">{cell_text}</td>')
                cell_index += colspan
            parts.append("</tr>")
        parts.append("</table>")
    return "\n".join(parts)


PHATHANH_TAG_STYLES = {
    "table": "width: 100%; border-collapse: collapse; margin-bottom: 20px; font-size: 14px;",
    "th": "border: 1px solid #bdc3c7; padding: 10px; text-align: left; vertical-align: middle;",
    "td": "border: 1px solid #bdc3c7; padding: 10px; text-align: left; vertical-align: middle;",
}


PHATHANH_CLASS_STYLES = {
    "container": "max-width: 800px; margin: 0 auto; border: 1px solid #e0e0e0; padding: 20px; border-radius: 8px;",
    "greeting": "font-size: 15px; margin-bottom: 20px;",
    "title-thanhly": "background-color: #FFC000; color: #000000; font-weight: bold; text-align: center; text-transform: uppercase; font-size: 14px; padding: 12px;",
    "title-khachhang": "background-color: #E2EFDA; color: #000000; font-weight: bold; text-align: center; text-transform: uppercase; font-size: 14px; padding: 10px;",
    "section-header": "background-color: #E2EFDA; font-weight: bold; text-align: center; color: #000000;",
    "col-header": "background-color: #ffffff; font-weight: bold; text-align: center; color: #000000;",
    "label-cell": "background-color: #ffffff; font-weight: bold; width: 130px;",
    "value-cell": "background-color: #ffffff;",
    "signature": "margin-top: 30px; border-top: 1px solid #e0e0e0; padding-top: 20px;",
    "sig-table": "border: none !important; width: 100%; border-collapse: collapse; background: none; margin: 0;",
    "sig-cell-logo": "border: none !important; width: 110px; vertical-align: middle; padding: 0 15px 0 0; background: none;",
    "sig-cell-info": "border: none !important; border-left: 2px solid #17b978 !important; padding-left: 15px; vertical-align: middle; background: none; line-height: 1.6; font-family: Arial, sans-serif;",
}


def _merge_inline_style(attrs: str, style: str) -> str:
    style_match = re.search(r'\sstyle=(["\'])(.*?)\1', attrs, flags=re.IGNORECASE | re.DOTALL)
    if style_match:
        merged = f"{style_match.group(2).rstrip(';')}; {style}"
        return attrs[:style_match.start()] + f' style="{merged}"' + attrs[style_match.end():]
    return f'{attrs} style="{style}"'


def inline_phathanh_email_styles(html_body: str) -> str:
    html_body = re.sub(r"<style\b[^>]*>.*?</style>", "", html_body, flags=re.IGNORECASE | re.DOTALL)

    for tag, style in PHATHANH_TAG_STYLES.items():
        pattern = re.compile(rf"<{tag}\b([^>]*)>", flags=re.IGNORECASE)
        html_body = pattern.sub(lambda match: f"<{tag}{_merge_inline_style(match.group(1), style)}>", html_body)

    class_pattern = re.compile(r"<([a-z0-9]+)\b([^>]*)\sclass=(['\"])(.*?)\3([^>]*)>", flags=re.IGNORECASE | re.DOTALL)

    def replace_class(match):
        tag, before, _quote, class_value, after = match.groups()
        styles = [
            PHATHANH_CLASS_STYLES[class_name]
            for class_name in str(class_value).split()
            if class_name in PHATHANH_CLASS_STYLES
        ]
        attrs = before + after
        if styles:
            attrs = _merge_inline_style(attrs, " ".join(styles))
        return f"<{tag}{attrs}>"

    html_body = class_pattern.sub(replace_class, html_body)
    body_match = re.search(r"<body\b[^>]*>(.*?)</body>", html_body, flags=re.IGNORECASE | re.DOTALL)
    if body_match:
        return body_match.group(1).strip()
    return html_body.strip()


def html_to_plain_text(html_content: str) -> str:
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        # Safe fallback if bs4 is missing
        text = re.sub(r'<br[^>]*>', '\n', html_content)
        text = re.sub(r'<[^>]+>', ' ', text)
        return re.sub(r' +', ' ', text).strip()

    soup = BeautifulSoup(html_content, "html.parser")
    parts = []
    
    def clean(text):
        return re.sub(r'\s+', ' ', text).strip()

    tables = soup.find_all("table")
    if not tables:
        return soup.get_text(separator="\n").strip()
        
    # 1. Process all <p> elements that are NOT inside a table as intro
    intro_text = []
    for p in soup.find_all("p"):
        if not p.find_parent("table"):
            txt = clean(p.get_text())
            if txt:
                intro_text.append(txt)
    if intro_text:
        parts.append("\n".join(intro_text))
        parts.append("")

    # 2. Process Tables
    for idx, table in enumerate(tables):
        table_text = clean(table.get_text())
        if idx > 0 and ("CENVALUE" in table_text or "Hotline" in table_text):
            sig_text = [
                "Trân trọng,",
                "CÔNG TY CỔ PHẦN THẨM ĐỊNH GIÁ THẾ KỶ (CENVALUE)",
                "VP TẠI GIA LAI: 90/60/3 Trường Chinh, phường Pleiku, tỉnh Gia Lai",
                "Hotline: 0905.226.968 | Website: www.cenvalue.vn"
            ]
            parts.append("\n".join(sig_text))
            continue
            
        rows = table.find_all("tr")
        if not rows:
            continue
            
        total_width = 104
        formatted_table = []
        
        def format_7_cols(cells):
            texts = [clean(c.get_text()) for c in cells]
            while len(texts) < 7:
                texts.append("")
            c0 = texts[0][:12]
            c1 = texts[1][:4]
            c2 = texts[2][:24]
            c3 = texts[3][:11]
            c4 = texts[4][:11]
            c5 = texts[5][:20]
            c6 = texts[6][:10]
            return (
                f" {c0:<12} | {c1:^4} | {c2:<24} | {c3:^11} | "
                f"{c4:<11} | {c5:<20} | {c6:^10}"
            )
            
        for r_idx, row in enumerate(rows):
            cells = row.find_all(["td", "th"])
            if not cells:
                continue
                
            colspans = [int(c.get("colspan", 1)) for c in cells]
            
            if len(cells) == 1 and colspans[0] >= 7:
                txt = clean(cells[0].get_text())
                border_char = "=" if "THANH LÝ" in txt or r_idx == 0 else "-"
                formatted_table.append(border_char * total_width)
                formatted_table.append(txt.center(total_width))
                formatted_table.append(border_char * total_width)
                
            elif len(cells) == 2 and colspans[0] == 2 and colspans[1] == 5:
                label = clean(cells[0].get_text())
                val = clean(cells[1].get_text())
                formatted_table.append(f" {label:<22} {val}")
                
            elif len(cells) == 2 and colspans[0] == 4 and colspans[1] == 3:
                text_left = clean(cells[0].get_text())
                text_right = clean(cells[1].get_text())
                if "Ghi chú" in text_left:
                    formatted_table.append("-" * total_width)
                    formatted_table.append(f" {text_left} {text_right}")
                    formatted_table.append("=" * total_width)
                else:
                    w_left = 68
                    w_right = total_width - w_left - 3
                    formatted_table.append("-" * total_width)
                    formatted_table.append(f" {text_left.center(w_left)} | {text_right.center(w_right)} ")
                    formatted_table.append("-" * total_width)
                    
            elif len(cells) == 7:
                line = format_7_cols(cells)
                formatted_table.append(line)
                if "Hình thức" in line:
                    formatted_table.append("-" * total_width)
            else:
                txts = [clean(c.get_text()) for c in cells]
                formatted_table.append(" | ".join(txts))
                
        parts.append("\n".join(formatted_table))
        
    return "\n\n".join([p for p in parts if p.strip()])


def build_phathanh_email_html(
    case: dict,
    recipient: str | None = None,
    include_signature: bool = True,
    logo_src: str = "cid:logo_cenvalue",
    delivery_quantity: str | int = "2",
) -> str:
    template_path = Path("src/templates/phathanh_template.html")
    if not template_path.exists():
        template_path = Path(__file__).parent / "templates" / "phathanh_template.html"
    if not template_path.exists():
        raise FileNotFoundError("Thi?u file template: src/templates/phathanh_template.html")

    html_template = template_path.read_text(encoding="utf-8")
    now = datetime.now()
    date_receive = (now + timedelta(days=1)).strftime("%d/%m/%Y")
    date_payment = now.strftime("%d/%m/%Y")
    signature = os.getenv("MAIL_SIGNATURE", "Trân trọng,<br>Century Appraisal")

    cccd_row = ""
    if case.get("customer_type") == "individual":
        citizen_id = case.get("citizen_id")
        if citizen_id:
            cccd_row = f'''
<tr>
<td colspan="2" class="label-cell">Số CCCD*:</td>
<td colspan="5" class="value-cell"><strong>{html.escape(str(citizen_id))}</strong></td>
</tr>'''

    quantity_text = str(delivery_quantity or "2").strip() or "2"
    recipient_html = format_phathanh_recipient_html(recipient)
    signature_block_html = ""
    if include_signature:
        signature_block_html = f'''
<!-- Restored Premium Signature layout as in user screenshot -->
<div class="signature">
<table class="sig-table">
<tr style="border: none !important; background: none;">
<td class="sig-cell-logo">
<img src="{logo_src}" alt="Cenvalue Logo" style="width: 100px; height: auto; display: block;" />
</td>
<td class="sig-cell-info">
<div style="font-size: 15px; font-weight: bold; color: #1e3d59;">CÔNG TY CỔ PHẦN THẨM ĐỊNH GIÁ THẾ KỶ (CENVALUE)</div>
<div style="font-size: 13px; color: #555555; margin-top: 4px;">
<strong>VP TẠI GIA LAI:</strong> 90/60/3 Trường Chinh, phường Pleiku, tỉnh Gia Lai
</div>
<div style="font-size: 13px; color: #555555; margin-top: 2px;">
<strong>Hotline:</strong> 0905.226.968 | <strong>Website:</strong> <a href="http://www.cenvalue.vn" target="_blank" style="color: #17b978; text-decoration: none; font-weight: bold;">www.cenvalue.vn</a>
</div>
</td>
</tr>
</table>
</div>'''

    replacements = {
        "{{ customer_name }}": html.escape(str(case.get("customer_info") or case.get("owner_name") or "")),
        "{{ customer_address }}": html.escape(str(case.get("customer_address") or case.get("dia_chi_thua_dat") or "")),
        "{{ cccd_row }}": cccd_row,
        "{{ delivery_quantity }}": html.escape(quantity_text),
        "{{ recipient_info }}": recipient_html,
        "{{ date_receive }}": date_receive,
        "{{ date_payment }}": date_payment,
        "{{ personal_note }}": "",
        "{{ email_signature }}": signature,
        "{{ signature_block }}": signature_block_html,
    }

    html_body = html_template
    for placeholder, value in replacements.items():
        html_body = html_body.replace(placeholder, str(value))

    # Sanitize line separators before inserting through browser automation.
    html_body = html_body.replace("\u2028", " ").replace("\u2029", " ").replace("\r", " ").replace("\n", " ")
    return inline_phathanh_email_styles(html_body)

def _decode_header_str(header_value):
    if not header_value:
        return ""
    try:
        decoded_parts = decode_header(str(header_value))
        parts = []
        for content, encoding in decoded_parts:
            if isinstance(content, bytes):
                parts.append(content.decode(encoding or 'utf-8', errors='replace'))
            else:
                parts.append(str(content))
        decoded = "".join(parts)
    except Exception:
        decoded = str(header_value)

    # Repair mojibake if any (e.g. UTF-8 interpreted as Latin-1 / CP1252)
    try:
        if any(c in decoded for c in "Æ°á»›º¡"):
            return decoded.encode('latin-1').decode('utf-8')
    except (UnicodeEncodeError, UnicodeDecodeError):
        pass
    try:
        if any(c in decoded for c in "Æ°á»›º¡"):
            return decoded.encode('cp1252').decode('utf-8')
    except (UnicodeEncodeError, UnicodeDecodeError):
        pass
    return decoded

def _find_latest_email_by_subject_sync(contract_number: str):
    settings = load_gmail_smtp_settings()
    # imap.gmail.com
    imap_host = os.getenv("IMAP_HOST", "imap.gmail.com")

    try:
        mail = imaplib.IMAP4_SSL(imap_host)
        mail.login(settings.username, settings.password)
        mail.select("inbox")

        # Search for contract number in subject
        search_query = f'SUBJECT "{contract_number}"'
        status, messages = mail.search(None, search_query)

        # Fallback: search for core part like "N06-0946" if full search yielded nothing
        if (status != "OK" or not messages[0]):
            import re as _re
            match = _re.search(r"([A-Za-z0-9]+-\d+)", contract_number)
            if match:
                core_part = match.group(1)
                search_query = f'SUBJECT "{core_part}"'
                status, messages = mail.search(None, search_query)

        if status != "OK" or not messages[0]:
            mail.logout()
            return None

        # Get latest message ID
        msg_ids = messages[0].split()
        latest_id = msg_ids[-1]

        # Fetch full email bytes and Gmail thread ID (X-GM-THRID)
        fetch_items = "(RFC822)"
        if imap_host == "imap.gmail.com":
            fetch_items = "(X-GM-THRID RFC822)"
        status, data = mail.fetch(latest_id, fetch_items)
        if status != "OK":
            mail.logout()
            return None

        # Extract Gmail thread ID and raw email content
        thread_id = ""
        raw_email = b""
        for part in data:
            if isinstance(part, tuple):
                if imap_host == "imap.gmail.com" and part[0]:
                    import re as _re
                    thrid_match = _re.search(rb'X-GM-THRID\s+(\d+)', part[0])
                    if thrid_match:
                        thread_id = format(int(thrid_match.group(1)), 'x')
                if len(part) > 1 and isinstance(part[1], bytes):
                    raw_email = part[1]
                    break

        from src.mail_listener import parse_incoming_email
        parsed = parse_incoming_email(raw_email)

        result = {
            "msg_id": _clean_header(parsed.message_id),
            "subject": _clean_header(parsed.subject),
            "from": _clean_header(parsed.from_email),
            "to": _clean_header(parsed.to_header),
            "cc": _clean_header(parsed.cc_header),
            "references": _clean_header(parsed.references),
            "thread_id": thread_id,
            "body": parsed.html if parsed.html else parsed.text,
        }

        mail.logout()
        return result
    except Exception as e:
        print(f"IMAP Error: {e}")
        return None

async def find_latest_email_by_subject(contract_number: str):
    from src.oauth2_service import get_enabled_oauth_provider, fetch_emails_via_oauth2
    provider = get_enabled_oauth_provider()
    if provider:
        try:
            print(f"[find_latest_email_by_subject] Searching via OAuth2 API for provider: {provider}")
            # Try searching with full contract number first
            oauth_emails = await fetch_emails_via_oauth2(
                provider, query_contract=contract_number, limit=5, unread_only=False
            )
            
            # Fallback to core part if no emails found
            if not oauth_emails and "-" in contract_number:
                import re as _re
                match = _re.search(r"([A-Za-z0-9]+-\d+)", contract_number)
                if match:
                    core_part = match.group(1)
                    print(f"[find_latest_email_by_subject] Fallback searching core part {core_part} via OAuth2")
                    oauth_emails = await fetch_emails_via_oauth2(
                        provider, query_contract=core_part, limit=5, unread_only=False
                    )
            
            if oauth_emails:
                latest_item = oauth_emails[0]
                raw_email = latest_item["raw_bytes"]
                
                from src.mail_listener import parse_incoming_email
                parsed = parse_incoming_email(raw_email)
                
                return {
                    "msg_id": _clean_header(parsed.message_id),
                    "subject": _clean_header(parsed.subject),
                    "from_email": _clean_header(parsed.from_email),
                    "to": _clean_header(parsed.to_header),
                    "cc": _clean_header(parsed.cc_header),
                    "references": _clean_header(parsed.references),
                    "thread_id": latest_item.get("thread_id", ""),
                    "body": parsed.html if parsed.html else parsed.text,
                }
        except Exception as exc:
            print(f"[find_latest_email_by_subject] OAuth2 search failed: {exc}")
            
    return await asyncio.to_thread(_find_latest_email_by_subject_sync, contract_number)

def _send_phathanh_reply_sync(original_mail, html_body, settings):
    msg = EmailMessage()

    # Subject: Re: Original Subject
    orig_subject = original_mail["subject"]
    if not orig_subject.lower().startswith("re:"):
        msg["Subject"] = f"Re: {orig_subject}"
    else:
        msg["Subject"] = orig_subject

    # Headers for threading
    msg["In-Reply-To"] = original_mail["msg_id"]
    references = original_mail["references"]
    msg["References"] = f"{references} {original_mail['msg_id']}".strip()

    # Recipients (Reply All logic)
    orig_from = original_mail["from"]
    orig_to = original_mail["to"] or ""
    orig_cc = original_mail["cc"] or ""

    all_to = _parse_email_list(orig_from)
    all_cc = _parse_email_list(orig_to) + _parse_email_list(orig_cc)

    # Deduplicate and remove self
    self_email = settings.username.lower()

    final_to = []
    seen = set()
    for addr in all_to:
        email_addr = parseaddr(addr)[1].lower()
        if email_addr and email_addr != self_email and email_addr not in seen:
            final_to.append(addr)
            seen.add(email_addr)

    final_cc = []
    for addr in all_cc:
        email_addr = parseaddr(addr)[1].lower()
        if email_addr and email_addr != self_email and email_addr not in seen:
            final_cc.append(addr)
            seen.add(email_addr)

    if not final_to:
        final_to = [orig_from] # Fallback

    msg["From"] = settings.mail_from
    msg["To"] = ", ".join(final_to)
    if final_cc:
        msg["Cc"] = ", ".join(final_cc)

    msg.set_content("Vui lòng xem nội dung trong định dạng HTML.")
    msg.add_alternative(html_body, subtype="html")

    # Inline image embedding (CID) for the logo
    parts = msg.get_payload()
    html_part = parts[-1]

    logo_path = Path(__file__).resolve().parent / "templates" / "logo.jpg"
    if logo_path.exists():
        import mimetypes
        ctype, encoding = mimetypes.guess_type(str(logo_path))
        if ctype is None or encoding is not None:
            ctype = 'image/jpeg'
        maintype, subtype = ctype.split('/', 1)

        try:
            with open(logo_path, 'rb') as f:
                img_data = f.read()
            html_part.add_related(img_data, maintype=maintype, subtype=subtype, cid='<logo_cenvalue>', filename='logo.jpg')
        except Exception as img_err:
            print(f"Error embedding logo image: {img_err}")

    # Send
    try:
        with smtplib.SMTP(settings.host, settings.port) as server:
            server.starttls()
            server.login(settings.username, settings.password)
            server.send_message(msg)
        return True
    except Exception as e:
        print(f"SMTP Error: {e}")
        return False

async def send_phathanh_reply(original_mail, html_body, settings):
    from src.oauth2_service import get_enabled_oauth_provider, send_email_via_oauth2
    from email.utils import parseaddr

    provider = get_enabled_oauth_provider()

    if provider:
        try:
            # Recipients (Reply All logic)
            orig_from = original_mail["from"]
            orig_to = original_mail["to"] or ""
            orig_cc = original_mail["cc"] or ""

            all_to = _parse_email_list(orig_from)
            all_cc = _parse_email_list(orig_to) + _parse_email_list(orig_cc)

            self_email = (settings.username or "").lower()

            final_to = []
            seen = set()
            for addr in all_to:
                email_addr = parseaddr(addr)[1].lower()
                if email_addr and email_addr != self_email and email_addr not in seen:
                    final_to.append(addr)
                    seen.add(email_addr)

            final_cc = []
            for addr in all_cc:
                email_addr = parseaddr(addr)[1].lower()
                if email_addr and email_addr != self_email and email_addr not in seen:
                    final_cc.append(addr)
                    seen.add(email_addr)

            if not final_to:
                final_to = [orig_from] # Fallback

            orig_subject = original_mail["subject"]
            subject = orig_subject if orig_subject.lower().startswith("re:") else f"Re: {orig_subject}"

            print(f"Replying to email using OAuth2 API for provider: {provider}")
            await send_email_via_oauth2(
                provider=provider,
                from_email=_clean_header(settings.mail_from or settings.username),
                to_email=", ".join(final_to),
                subject=subject,
                html_body=html_body,
                cc_emails=final_cc,
                reply_to_msg_id=original_mail["msg_id"],
                references=original_mail["references"],
                thread_id=original_mail.get("thread_id") or None,
            )
            return True
        except Exception as exc:
            raise RuntimeError(f"Gửi mail phát hành qua {provider.upper()} OAuth2 thất bại: {exc}") from exc

    return await asyncio.to_thread(_send_phathanh_reply_sync, original_mail, html_body, settings)


async def send_phathanh_email_for_case(case: dict, recipient: str = None) -> str:
    """Gửi email phát hành chứng thư cho một hồ sơ."""
    contract_number = case.get("contract_number")
    if not contract_number:
        raise ValueError("Hồ sơ không có số hợp đồng.")

    # Auto-lookup customer data from DB if missing
    if not case.get("customer_info"):
        try:
            from src.database_manager import get_case_by_contract_number
            cases_db = os.getenv("TELEGRAM_CASES_DB", os.getenv("SQLITE_DATABASE", "data/cases.db"))
            # Try multiple search terms: full contract, without trailing suffix, hyphenated part
            search_terms = [contract_number]
            parts = [p.strip() for p in contract_number.split('/') if p.strip()]
            # Try without last part if it looks like a suffix (e.g. DN, DT)
            if len(parts) > 1 and len(parts[-1]) <= 3:
                search_terms.append('/'.join(parts[:-1]))
            # Try hyphenated part (e.g. D10-0104)
            for p in parts:
                if '-' in p and len(p) >= 5:
                    search_terms.append(p)
                    break
            db_record = None
            for term in search_terms:
                db_record = await get_case_by_contract_number(cases_db, term)
                if db_record:
                    break
            if db_record:
                for k, v in db_record.items():
                    if v and k not in case:
                        case[k] = v
                print(f"[send_phathanh] Auto-loaded customer data from DB: {case.get('customer_info', 'N/A')}")
            else:
                print(f"[send_phathanh] DB auto-lookup: no record found for {search_terms}")
        except Exception as db_exc:
            print(f"[send_phathanh] DB auto-lookup skipped: {db_exc}")

    include_signature = phathanh_send_mode() != "playwright"
    html_body = build_phathanh_email_html(case, recipient, include_signature=include_signature)
    if phathanh_send_mode() == "playwright":
        return await send_phathanh_email_via_playwright_raw(case, html_body)

    from src.oauth2_service import get_enabled_oauth_provider, fetch_emails_via_oauth2, send_email_via_oauth2

    provider = get_enabled_oauth_provider()

    original_mail = None
    if provider:
        try:
            print(f"Finding original email for {contract_number} via OAuth2 API for provider: {provider}")
            # Search ALL emails (not just unread) to find the original thread
            oauth_emails = await fetch_emails_via_oauth2(
                provider, query_contract=contract_number, limit=5, unread_only=False
            )
            if oauth_emails:
                # Get latest raw email
                latest_item = oauth_emails[0]
                raw_email = latest_item["raw_bytes"]
                msg = email.message_from_bytes(raw_email)
                original_mail = {
                    "msg_id": _clean_header(msg.get("Message-ID")),
                    "subject": _clean_header(_decode_header_str(msg.get("Subject"))),
                    "from": _clean_header(_decode_header_str(msg.get("From"))),
                    "to": _clean_header(_decode_header_str(msg.get("To"))),
                    "cc": _clean_header(_decode_header_str(msg.get("Cc"))),
                    "references": _clean_header(msg.get("References", "")),
                    "thread_id": latest_item.get("thread_id", ""),
                }
                print(f"Found original email: subject='{original_mail['subject']}', thread_id='{original_mail.get('thread_id', '')}', msg_id='{original_mail['msg_id']}'")
        except Exception as exc:
            raise RuntimeError(f"Đọc mail gốc qua {provider.upper()} OAuth2 thất bại: {exc}") from exc

    if provider and not original_mail:
        raise RuntimeError(f"Không tìm thấy email nào trên {provider.upper()} có tiêu đề chứa số hợp đồng: {contract_number}.")
    if not provider:
        # 1. Tìm mail cũ bằng IMAP (truyền thống)
        original_mail = await find_latest_email_by_subject(contract_number)
        if not original_mail:
            raise RuntimeError(f"Không tìm thấy email nào có tiêu đề chứa số hợp đồng: {contract_number}. Vui lòng gửi email yêu cầu định giá trước.")

    # html_body was already built by build_phathanh_email_html() above.
    # It handles both DOCX template and HTML template fallback.

    # 3. Gửi mail
    settings = load_gmail_smtp_settings()

    success = False
    if provider:
        try:
            # Recipients (Reply All logic)
            orig_from = original_mail["from"]
            orig_to = original_mail["to"] or ""
            orig_cc = original_mail["cc"] or ""

            all_to = _parse_email_list(orig_from)
            all_cc = _parse_email_list(orig_to) + _parse_email_list(orig_cc)

            self_email = (settings.username or "").lower()

            final_to = []
            seen = set()
            for addr in all_to:
                email_addr = parseaddr(addr)[1].lower()
                if email_addr and email_addr != self_email and email_addr not in seen:
                    final_to.append(addr)
                    seen.add(email_addr)

            final_cc = []
            for addr in all_cc:
                email_addr = parseaddr(addr)[1].lower()
                if email_addr and email_addr != self_email and email_addr not in seen:
                    final_cc.append(addr)
                    seen.add(email_addr)

            if not final_to:
                final_to = [orig_from] # Fallback

            orig_subject = original_mail["subject"]
            subject = orig_subject if orig_subject.lower().startswith("re:") else f"Re: {orig_subject}"

            thread_id = original_mail.get("thread_id") or None
            print(f"Replying to email using OAuth2 API for provider: {provider}, thread_id={thread_id}")
            await send_email_via_oauth2(
                provider=provider,
                from_email=_clean_header(settings.mail_from or settings.username),
                to_email=", ".join(final_to),
                subject=subject,
                html_body=html_body,
                cc_emails=final_cc,
                reply_to_msg_id=original_mail["msg_id"],
                references=original_mail["references"],
                thread_id=thread_id,
            )
            success = True
        except Exception as exc:
            raise RuntimeError(f"Gửi mail phát hành qua {provider.upper()} OAuth2 thất bại: {exc}") from exc

    if not success:
        success = await send_phathanh_reply(original_mail, html_body, settings)
        if not success:
            raise RuntimeError("Gửi email SMTP thất bại.")

    return original_mail.get("to") or original_mail.get("from")


async def send_phathanh_email_via_browser_use(case_data: dict, html_body: str) -> str:
    """Gửi email phát hành chứng thư thông qua Webmail (OWA/Outlook) bằng Browser Use và Gemini API."""
    import os
    import asyncio
    from browser_use import Agent, Browser, ChatGoogle
    api_keys = gemini_keys_with_backups()
    model = os.getenv("GEMINI_MODEL", "gemini-2.5-flash").strip() or "gemini-2.5-flash"
    username = os.getenv("MAIL_USERNAME")
    password = os.getenv("MAIL_PASSWORD")
    user_data_dir = os.getenv("CHROME_USER_DATA_DIR")
    webmail_url = os.getenv("WEBMAIL_URL", "https://owa.cengroup.vn/").strip()

    contract_number = str(case_data.get("contract_number") or "").strip()
    if not contract_number:
        raise ValueError("Hồ sơ không có số hợp đồng để tìm kiếm email.")

    if not api_keys:
        raise RuntimeError("Thiếu biến môi trường GEMINI_API_KEY để chạy AI Agent.")

    llm = ChatGoogle(
        model=model,
        api_key=api_keys[0]
    )
    fallback_llm = None
    if len(api_keys) > 1:
        fallback_llm = ChatGoogle(
            model=model,
            api_key=api_keys[1]
        )

    # Khởi tạo Browser ẩn danh (headless=True trên Linux/VPS, False trên Windows để test trực quan)
    is_headless = os.getenv("PLAYWRIGHT_HEADLESS", "true" if os.name != "nt" else "false").strip().casefold() in ("true", "1", "yes")

    # Chrome arguments required for Linux root context execution
    chrome_args = ["--no-sandbox", "--disable-dev-shm-usage", "--disable-gpu"]

    browser_kwargs = {
        "headless": is_headless,
        "args": chrome_args,
        "enable_default_extensions": False
    }
    if user_data_dir and os.path.exists(user_data_dir):
        browser_kwargs["user_data_dir"] = user_data_dir

    browser = Browser(**browser_kwargs)
    tools = build_phathanh_browser_tools(html_body)

    # Prompt chỉ dẫn chi tiết cho AI Agent điều khiển OWA
    prompt = (
        f"Hãy thực hiện các bước sau để gửi email phát hành chứng thư:\n"
        f"1. Truy cập vào trang Webmail {webmail_url}.\n"
        f"2. Nếu chưa đăng nhập, sử dụng tài khoản '{username}' và mật khẩu '{password}' để đăng nhập.\n"
        f"3. Nhập từ khóa số hợp đồng '{contract_number}' vào ô tìm kiếm (Search) ở trên cùng và nhấn Enter.\n"
        f"4. Click mở email mới nhất trong kết quả tìm kiếm.\n"
        f"5. Tìm và click nút 'Trả lời tất cả' (Reply All).\n"
        f"6. Tìm ô soạn thảo nội dung trả lời (thường là ô nhập liệu contenteditable='true' hoặc role='textbox').\n"
        f"7. Nhập nội dung sau vào ô soạn thảo: 'Dear all, Phòng thẩm định cho phát hành chứng thư giúp mình nhé, mình cảm ơn!' "
        f"và dán bảng thông tin này vào. Để dán bảng thông tin một cách chuẩn xác nhất, hãy thực hiện chạy JavaScript (execute_javascript) để đặt `innerHTML` của ô soạn thảo bằng đoạn mã HTML sau:\n"
        f"\"\"\"{html_body}\"\"\"\n"
        f"8. Sau khi đã chèn thành công nội dung và bảng thông tin xuất hiện chính xác trên giao diện, hãy click nút 'Gửi' (Send) để hoàn thành gửi mail."
    )

    prompt = (
        f"Use OWA Webmail at {webmail_url} to send the certificate release reply.\n"
        f"1. Log in with the username '{username}' and password '{password}' if needed.\n"
        f"2. Search for the email whose subject contains this contract number: {contract_number}.\n"
        f"3. Click on the first/newest email item in the search results list to open its details. Wait for the email details pane to load completely.\n"
        f"4. Look for and click the 'Reply All' button (or 'Trả lời tất cả' in Vietnamese) inside the opened email pane.\n"
        f"5. Once the reply editor opens, click the editor body to focus it, then call the `insert_phathanh_html` tool exactly once to insert the prepared release content.\n"
        f"6. Confirm the content is inserted correctly, then click the 'Send' button (or 'Gửi' in Vietnamese) to send the email."
    )

    try:
        agent = None
        try:
            agent = Agent(
                task=prompt,
                llm=llm,
                fallback_llm=fallback_llm,
                browser=browser,
                tools=tools
            )
            history = await agent.run()
            if not history or not history.is_successful():
                raise RuntimeError("AI Agent execution was not successful (likely rate limited or failed).")
        except Exception as e:
            # Catch inner exceptions (like rate limits / API errors during run)
            raise e
        finally:
            try:
                await browser.stop()
            except Exception as e:
                import logging
                logging.getLogger(__name__).warning(f"Failed to stop browser: {e}")
    except Exception as agent_exc:
        import logging
        logging.getLogger(__name__).error(f"Browser Use agent failed: {agent_exc}. Falling back to raw Playwright flow...")
        try:
            # Ensure the browser session is fully stopped before starting fallback
            await browser.stop()
        except Exception:
            pass
        return await send_phathanh_email_via_playwright_raw(case_data, html_body)

    return case_data.get("to") or case_data.get("from") or "Webmail AI Agent"


async def send_phathanh_email_via_playwright_raw(case_data: dict, html_body: str) -> str:
    """Fallback method using raw Playwright commands to reply to the certificate release email."""
    import os
    import asyncio
    from playwright.async_api import async_playwright

    username = os.getenv("MAIL_USERNAME")
    password = os.getenv("MAIL_PASSWORD")
    user_data_dir = os.getenv("CHROME_USER_DATA_DIR")
    if not user_data_dir:
        # Default to a local persistent directory to avoid logon session limits and reuse session
        user_data_dir = os.path.join(os.path.expanduser("~"), ".gemini", "antigravity", "chrome_profile")

    webmail_url = os.getenv("WEBMAIL_URL", "https://owa.cengroup.vn/").strip()
    is_headless = os.getenv("PLAYWRIGHT_HEADLESS", "true" if os.name != "nt" else "false").strip().casefold() in ("true", "1", "yes")

    contract_number = str(case_data.get("contract_number") or "").strip()
    if not contract_number:
        raise ValueError("Hồ sơ không có số hợp đồng để tìm kiếm email.")

    chrome_args = ["--no-sandbox", "--disable-dev-shm-usage", "--disable-gpu"]

    print(f"[Playwright Raw Fallback] Starting raw Playwright flow for contract {contract_number}")

    user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"

    async with async_playwright() as p:
        if user_data_dir:
            os.makedirs(user_data_dir, exist_ok=True)
            context = await p.chromium.launch_persistent_context(
                user_data_dir=user_data_dir,
                headless=is_headless,
                user_agent=user_agent,
                args=chrome_args,
            )
            page = context.pages[0] if context.pages else await context.new_page()
        else:
            browser = await p.chromium.launch(
                headless=is_headless,
                args=chrome_args,
            )
            context = await browser.new_context()
            page = await context.new_page()

        try:
            # 1. Navigate to the login page
            print(f"[Playwright Raw Fallback] Navigating to {webmail_url}")
            await page.goto(webmail_url, timeout=60000)
            await page.wait_for_load_state("load")

            # Check if we are already logged in
            username_selector = "input#username, input[name='username'], input[type='email'], input[type='text']"
            try:
                await page.wait_for_selector(username_selector, timeout=5000)
                needs_login = True
            except Exception:
                needs_login = False

            if needs_login:
                print("[Playwright Raw Fallback] Login form detected. Logging in...")
                await page.fill("input#username, input[name='username'], input[type='email']", username)
                await page.fill("input#password, input[name='password'], input[type='password']", password)
                await page.keyboard.press("Enter")
                print("[Playwright Raw Fallback] Submitted login form via Enter. Waiting for navigation/dashboard...")
                await page.wait_for_load_state("networkidle", timeout=15000)
                await page.wait_for_timeout(10000)
            else:
                print("[Playwright Raw Fallback] Already logged in or login form not visible.")
                await page.wait_for_timeout(5000)

            # Handle OWA error pages (e.g. "too many active sessions")
            for _retry in range(3):
                error_page = await page.evaluate(
                    r"""
                    () => {
                      const text = document.body ? (document.body.innerText || '') : '';
                      if (/Something went wrong|too many active sessions|Refresh the page/i.test(text)) {
                        const refreshLink = document.querySelector('a');
                        if (refreshLink && /refresh/i.test(refreshLink.textContent || '')) {
                          refreshLink.click();
                          return { error: true, action: 'clicked_refresh' };
                        }
                        return { error: true, action: 'detected_but_no_refresh_link' };
                      }
                      return { error: false };
                    }
                    """
                )
                if error_page and error_page.get("error"):
                    print(f"[Playwright Raw Fallback] OWA error page detected: {error_page}. Refreshing...")
                    await page.reload(wait_until="load", timeout=30000)
                    await page.wait_for_timeout(10000)
                else:
                    break

            # OWA first-run setup can appear after login on a fresh browser profile.
            try:
                first_run_done = await page.evaluate(
                r"""
                    () => {
                      const text = document.body ? (document.body.innerText || '') : '';
                      const selects = Array.from(document.querySelectorAll('select'));
                      if (!selects.length || !/Outlook|m\u00fai gi\u1edd|time zone|ng\u00f4n ng\u1eef|language/i.test(text)) {
                        return { handled: false, reason: 'first-run setup not visible' };
                      }
                      const timezoneSelect = selects[selects.length - 1];
                      const options = Array.from(timezoneSelect.options || []);
                      const preferred = options.find((option) => /Bangkok|Hanoi|Ha Noi|Jakarta|\\+07|UTC\\+7|SE Asia/i.test(option.textContent || option.value));
                      const fallback = options.find((option) => option.value && !/ch\u1ecdn|select/i.test(option.textContent || ''));
                      const chosen = preferred || fallback;
                      if (chosen) {
                        timezoneSelect.value = chosen.value;
                        timezoneSelect.dispatchEvent(new Event('change', { bubbles: true }));
                      }
                      const buttons = Array.from(document.querySelectorAll('button, input[type="submit"], [role="button"], div[role="button"]'));
                      const saveButton = buttons.find((el) => /save|l\u01b0u/i.test(el.innerText || el.textContent || el.value || el.getAttribute('aria-label') || ''));
                      if (!saveButton) return { handled: false, reason: 'save button not found' };
                      saveButton.click();
                      return { handled: true, timezone: chosen ? (chosen.textContent || chosen.value) : null };
                     }
                     """
                )
                if first_run_done and first_run_done.get("handled"):
                    print(f"[Playwright Raw Fallback] Completed OWA first-run setup: {first_run_done}")
                    await page.wait_for_load_state("networkidle", timeout=20000)
                    await page.wait_for_timeout(5000)
            except Exception as setup_exc:
                print(f"[Playwright Raw Fallback] OWA first-run setup check skipped: {setup_exc}")

            # 2. Focus OWA search input via Alt+Q shortcut.
            print("[Playwright Raw Fallback] Focusing OWA search input via Alt+Q...")
            await page.keyboard.press("Alt+q")
            await page.wait_for_timeout(2000)

            # Verify search input is focused
            search_focus = None
            for _ in range(10):
                search_focus = await page.evaluate(
                    r"""
                    () => {
                      const isVisible = (el) => {
                        const rect = el.getBoundingClientRect();
                        return rect.width > 0 && rect.height > 0;
                      };
                      const inputs = Array.from(document.querySelectorAll('input')).filter(isVisible);
                      const searchInput = inputs.find(i => {
                        const h = [i.type, i.title, i.placeholder, i.getAttribute('aria-label'), i.id].join(' ').toLowerCase();
                        return h.includes('search') || h.includes('tìm kiếm');
                      }) || inputs.find(i => i.type === 'text');
                      if (searchInput) {
                        searchInput.focus();
                        searchInput.click();
                        const rect = searchInput.getBoundingClientRect();
                        return { ok: true, rect: { x: rect.x, y: rect.y, w: rect.width, h: rect.height } };
                      }
                      return { ok: false };
                    }
                    """
                )
                if search_focus and search_focus.get("ok"):
                    break
                await page.wait_for_timeout(1000)

            if not search_focus or not search_focus.get("ok"):
                raise RuntimeError(f"OWA search input not found after Alt+Q: {search_focus}")
            print(f"[Playwright Raw Fallback] Search input focused: {search_focus}")

            # 3. Input contract number and run search.
            # Extract a cleaner search query if the contract number contains slashes (e.g., D10-0104 from 010/2024/D10-0104/DN)
            search_query = contract_number
            parts = [p.strip() for p in contract_number.split('/') if p.strip()]
            for p in parts:
                if '-' in p and len(p) >= 5:
                    search_query = p
                    break
            else:
                for p in parts:
                    if not p.isdigit() and p.upper() != 'DN' and len(p) >= 4:
                        search_query = p
                        break

            print(f"[Playwright Raw Fallback] Inputting search query: {search_query} (derived from contract: {contract_number})")
            await page.keyboard.press("Control+A")
            await page.keyboard.press("Delete")
            await page.keyboard.type(search_query, delay=20)
            search_clicked = await page.evaluate(
                r"""
                () => {
                  const isVisible = (el) => {
                    const rect = el.getBoundingClientRect();
                    const style = window.getComputedStyle(el);
                    return rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none';
                  };
                  const normalize = (value) => (value || '').toLocaleLowerCase('vi-VN');
                  const buttons = Array.from(document.querySelectorAll('button, a, [role="button"], input[type="submit"]'));
                  const searchButton = buttons.find((el) => {
                    if (!isVisible(el)) return false;
                    const text = normalize([el.title, el.getAttribute('aria-label'), el.innerText, el.textContent, el.value].filter(Boolean).join(' '));
                    const rect = el.getBoundingClientRect();
                    return (text.includes('search') || text.includes('tìm kiếm')) && rect.y < 80;
                  });
                  if (!searchButton) return false;
                  searchButton.click();
                  return true;
                }
                """
            )
            if search_clicked:
                print("[Playwright Raw Fallback] Search submitted via search button.")
            else:
                await page.keyboard.press("Enter")
                print("[Playwright Raw Fallback] Search submitted via Enter.")

            print("[Playwright Raw Fallback] Waiting search results to load...")
            await page.wait_for_timeout(8000)

            opened_result = await page.evaluate(
                r"""
                (contractNumber) => {
                  const normalize = (value) => (value || '').replace(/\s+/g, ' ').trim();
                  const contract = normalize(contractNumber);
                  const parts = contract.split('/');
                  const hyphenPart = parts.find(p => p.includes('-') && p.length >= 5) || '';
                  const tail = parts.filter(Boolean).pop() || contract;
                  const tailPrefix = tail.slice(0, Math.max(5, Math.min(tail.length, 7)));
                  const terms = Array.from(new Set([contract, hyphenPart, tail, tailPrefix].filter((term) => term && term.length >= 3)));
                  const isVisible = (el) => {
                    const rect = el.getBoundingClientRect();
                    const style = window.getComputedStyle(el);
                    return rect.width > 30 && rect.height > 10 && style.visibility !== 'hidden' && style.display !== 'none';
                  };
                  const clickTargetFor = (el) => el.closest('a, [role="option"], [role="listitem"], [role="button"], button, tr') || el;
                  const nodes = Array.from(document.querySelectorAll('tr, a, td, div, span'));
                  const mapped = nodes
                    .map((el) => ({ el, target: clickTargetFor(el), text: normalize(el.innerText || el.textContent), rect: el.getBoundingClientRect(), tag: el.tagName }))
                    .filter((item) => isVisible(item.el));
                  const rowLike = (item) => item.rect.y > 120 && item.rect.height >= 12 && item.rect.height < 80 && item.rect.width > 80;
                  const textMatches = mapped
                    .filter((item) => rowLike(item) && terms.some((term) => item.text.includes(term)))
                    .sort((a, b) => {
                      const aScore = (a.text.includes(contract) ? 0 : a.text.includes(hyphenPart) ? 1 : a.text.includes(tail) ? 2 : 3) + (a.tag === 'A' ? 0 : a.tag === 'TR' ? 1 : 2);
                      const bScore = (b.text.includes(contract) ? 0 : b.text.includes(hyphenPart) ? 1 : b.text.includes(tail) ? 2 : 3) + (b.tag === 'A' ? 0 : b.tag === 'TR' ? 1 : 2);
                      return (aScore - bScore) || (a.rect.y - b.rect.y) || (a.rect.x - b.rect.x);
                    });
                  const item = textMatches[0];
                  if (!item) return { ok: false, reason: 'no matching search result found', terms };
                  item.target.click();
                  return { ok: true, text: item.text.slice(0, 160), rect: { x: item.rect.x, y: item.rect.y, w: item.rect.width, h: item.rect.height }, tag: item.tag, terms };
                }
                """,
                contract_number,
            )
            if opened_result and opened_result.get("ok"):
                print(f"[Playwright Raw Fallback] Opened matching search result: {opened_result}")
            else:
                raise RuntimeError(f"Could not find or open matching search result: {opened_result}")

            await page.wait_for_timeout(4000)
            # 4. Open Reply All composer. Classic OWA supports a direct ReplyAll action URL.
            classic_reply_opened = False
            try:
                current_url = page.url
                if "ae=Item" in current_url and "id=" in current_url:
                    from urllib.parse import parse_qs, urlencode, urlparse, urlunparse
                    parsed = urlparse(current_url)
                    query = parse_qs(parsed.query)
                    message_id = (query.get("id") or [None])[0]
                    if message_id:
                        reply_query = {
                            "ae": "PreFormAction",
                            "a": "ReplyAll",
                            "t": query.get("t", ["IPM.Note"])[0],
                            "id": message_id,
                        }
                        reply_url = urlunparse((parsed.scheme, parsed.netloc, parsed.path, "", urlencode(reply_query), ""))
                        print(f"[Playwright Raw Fallback] Opening classic OWA ReplyAll URL: {reply_url}")
                        await page.goto(reply_url, wait_until="domcontentloaded", timeout=60000)
                        await page.wait_for_timeout(5000)
                        classic_reply_opened = True
            except Exception as reply_url_exc:
                print(f"[Playwright Raw Fallback] Classic ReplyAll URL fallback skipped: {reply_url_exc}")

            # 4b. Click Reply All if direct composer URL was not available.
            print("[Playwright Raw Fallback] Locating Reply All button...")
            reply_all_result = {"ok": True, "text": "classic ReplyAll URL"} if classic_reply_opened else await page.evaluate(
                r"""
                () => {
                  const normalize = (value) => (value || '').toLocaleLowerCase('vi-VN').normalize('NFD').replace(/[\u0300-\u036f]/g, '').replace(/\\s+/g, ' ').trim();
                  const replyAllTerms = [
                    'reply all',
                    'tra loi tat ca',
                    'tra loi tat',
                    'tr\u1ea3 l\u1eddi t\u1ea5t c\u1ea3',
                    'tr\u1ea3 l\u1eddi t\u1ea5t'
                  ];
                  const hasReplyAll = (text) => {
                    const value = normalize(text);
                    return replyAllTerms.some((term) => value.includes(term));
                  };
                  const isVisible = (el) => {
                    const rect = el.getBoundingClientRect();
                    const style = window.getComputedStyle(el);
                    return rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none';
                  };
                  const nodes = Array.from(document.querySelectorAll('button, a, [role="button"], [role="menuitem"]'));
                  const candidates = nodes
                    .map((el) => ({
                      el,
                      text: [el.getAttribute('aria-label'), el.getAttribute('title'), el.innerText, el.textContent].filter(Boolean).join(' '),
                      rect: el.getBoundingClientRect(),
                    }))
                    .filter((item) => isVisible(item.el) && hasReplyAll(item.text))
                    .sort((a, b) => (a.rect.y - b.rect.y) || (a.rect.x - b.rect.x));
                  const item = candidates[0];
                  if (!item) return { ok: false, reason: 'reply all button not found' };
                  const target = item.el.closest('button, [role="button"]') || item.el;
                  target.click();
                  return { ok: true, text: item.text.slice(0, 160) };
                }
                """
            )
            reply_all_clicked = bool(reply_all_result and reply_all_result.get("ok"))
            if reply_all_clicked:
                print(f"[Playwright Raw Fallback] Reply All clicked via JS: {reply_all_result}")
            else:
                print(f"[Playwright Raw Fallback] Reply All JS lookup failed: {reply_all_result}")
                for selector in (
                    "button:has-text('Reply All')",
                    "button:has-text('Tr\u1ea3 l\u1eddi t\u1ea5t c\u1ea3')",
                    "[aria-label*='Reply all']",
                    "[aria-label*='Reply All']",
                    "[aria-label*='Tr\u1ea3 l\u1eddi t\u1ea5t']",
                    "[title*='Reply all']",
                    "[title*='Reply All']",
                    "[title*='Tr\u1ea3 l\u1eddi t\u1ea5t']",
                    "div[role='button']:has-text('Reply All')",
                    "span:has-text('Reply All')",
                ):
                    try:
                        target = page.locator(selector).first
                        if await target.count() > 0:
                            await target.click(timeout=5000)
                            reply_all_clicked = True
                            print(f"[Playwright Raw Fallback] Reply All clicked via selector: {selector}")
                            break
                    except Exception:
                        continue

            if not reply_all_clicked:
                print("[Playwright Raw Fallback] Reply All button not found. Trying Ctrl+Shift+R shortcut...")
                await page.keyboard.press("Control+Shift+R")
                await page.wait_for_timeout(3000)
                reply_editor_opened = await page.evaluate(
                r"""
                    () => Array.from(document.querySelectorAll('[contenteditable="true"], [role="textbox"], [aria-label*="Message body"], [aria-label*="body"]'))
                      .some((el) => {
                        const rect = el.getBoundingClientRect();
                        const style = window.getComputedStyle(el);
                        return rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none';
                      })
                    """
                )
                if not reply_editor_opened:
                    raise RuntimeError("Reply All button not found in OWA and Ctrl+Shift+R did not open reply editor.")
                print("[Playwright Raw Fallback] Reply editor opened via Ctrl+Shift+R fallback.")

            await page.wait_for_timeout(3000)

            # 5. Use JS insert html_body into active editor. Classic OWA uses an iframe editor.
            print("[Playwright Raw Fallback] Waiting for visible reply editor...")
            editor_found = False
            for _ in range(30):
                editor_found = await page.evaluate(
                    r"""
                    () => {
                      const isVisible = (el) => {
                        const rect = el.getBoundingClientRect();
                        const style = window.getComputedStyle(el);
                        return rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none';
                      };
                      return Array.from(document.querySelectorAll("iframe, textarea, [contenteditable='true'], [role='textbox'], [aria-label*='Message body'], [aria-label*='body']")).some(isVisible);
                    }
                    """
                )
                if editor_found:
                    break
                await page.wait_for_timeout(1000)
            if not editor_found:
                raise RuntimeError("Reply editor did not appear (no visible matching elements found).")

            print("[Playwright Raw Fallback] Evaluating JS insert HTML into editor...")
            plain_body = html_to_plain_text(html_body)
            result = await page.evaluate(
                r"""
                async (args) => {
                  const html = args.html;
                  const plain = args.plain;
                  const isVisible = (el) => {
                    const rect = el.getBoundingClientRect();
                    const style = window.getComputedStyle(el);
                    return rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none';
                  };
                  const visible = (selector) => Array.from(document.querySelectorAll(selector)).filter(isVisible);
                  let editor = null;
                  let mode = null;

                  const iframes = visible('iframe');
                  for (const frame of iframes) {
                    try {
                      const doc = frame.contentDocument || frame.contentWindow.document;
                      if (doc && doc.body) {
                        editor = doc.body;
                        mode = 'iframe';
                        break;
                      }
                    } catch (error) {}
                  }

                  if (!editor) {
                    const textareas = visible('textarea');
                    if (textareas.length) {
                      editor = textareas[textareas.length - 1];
                      mode = 'textarea';
                    }
                  }

                  if (!editor) {
                    const candidates = visible('[contenteditable="true"], [role="textbox"], [aria-label*="Message body"], [aria-label*="body"]');
                    editor = candidates[candidates.length - 1];
                    mode = 'contenteditable';
                  }

                  if (!editor) return { ok: false, reason: 'reply editor not found' };

                  if (mode === 'textarea') {
                    editor.focus();
                    editor.value = plain;
                    editor.dispatchEvent(new Event('input', { bubbles: true }));
                    editor.dispatchEvent(new Event('change', { bubbles: true }));
                    return { ok: true, mode, text: editor.value.slice(0, 160) };
                  }

                  editor.focus();
                  editor.innerHTML = html;
                  editor.dispatchEvent(new InputEvent('input', { bubbles: true, inputType: 'insertHTML', data: html }));
                  editor.dispatchEvent(new Event('change', { bubbles: true }));
                  const renderedText = (editor.innerText || editor.textContent || '').trim();
                  return {
                    ok: renderedText.includes('Dear all') || editor.innerHTML.includes('Dear all'),
                    mode,
                    text: renderedText.slice(0, 160)
                  };
                }
                """,
                {"html": html_body, "plain": plain_body},
            )
            print(f"[Playwright Raw Fallback] Editor insertion result: {result}")
            if not result or not result.get("ok"):
                raise RuntimeError(f"Failed to insert phathanh content into OWA editor: {result.get('reason') if result else 'unknown error'}")

            # 6. Send email via Ctrl+Enter
            print("[Playwright Raw Fallback] Sending email via Ctrl+Enter...")
            await page.evaluate(
                r"""
                () => {
                  const editors = Array.from(document.querySelectorAll(
                    '[contenteditable="true"], [role="textbox"][contenteditable="true"]'
                  )).filter(el => {
                    const r = el.getBoundingClientRect();
                    return r.width > 200 && r.height > 50;
                  });
                  if (editors.length) editors[editors.length - 1].focus();
                }
                """
            )
            await page.wait_for_timeout(500)
            await page.keyboard.press("Control+Enter")
            print("[Playwright Raw Fallback] Ctrl+Enter sent.")

            # 7. Wait for compose editor to close + Undo Send progress bar to finish
            print("[Playwright Raw Fallback] Waiting for editor to close...")
            try:
                await page.wait_for_selector('[contenteditable="true"]', state="hidden", timeout=10000)
                print("[Playwright Raw Fallback] Editor closed.")
            except Exception:
                print("[Playwright Raw Fallback] Editor did not close within 10s.")

            # Wait for "Hủy bỏ gửi" (Undo Send) progress bar to appear and then disappear
            print("[Playwright Raw Fallback] Checking for OWA 'Hủy bỏ gửi' (Undo Send) progress bar...")
            has_undo = False
            for _ in range(6):
                undo_exists = await page.evaluate(
                    r"""
                    () => {
                      const normalize = (val) => (val || '').trim().toLowerCase();
                      const els = Array.from(document.querySelectorAll('span, div, button, a'));
                      return els.some(el => {
                        const text = normalize(el.innerText || el.textContent || '');
                        return text.includes('hủy bỏ gửi') || text.includes('hủy gửi') || text.includes('undo send');
                      });
                    }
                    """
                )
                if undo_exists:
                    has_undo = True
                    print("[Playwright Raw Fallback] 'Hủy bỏ gửi' progress bar detected. Waiting for it to complete...")
                    break
                await page.wait_for_timeout(500)

            if has_undo:
                # Wait for it to disappear (max 30s)
                for _ in range(60):
                    undo_exists = await page.evaluate(
                        r"""
                        () => {
                          const normalize = (val) => (val || '').trim().toLowerCase();
                          const els = Array.from(document.querySelectorAll('span, div, button, a'));
                          return els.some(el => {
                            const text = normalize(el.innerText || el.textContent || '');
                            const rect = el.getBoundingClientRect();
                            return (text.includes('hủy bỏ gửi') || text.includes('hủy gửi') || text.includes('undo send')) && rect.width > 0 && rect.height > 0;
                          });
                        }
                        """
                    )
                    if not undo_exists:
                        print("[Playwright Raw Fallback] 'Hủy bỏ gửi' progress bar completed. Mail is sent!")
                        break
                    await page.wait_for_timeout(500)
            else:
                print("[Playwright Raw Fallback] No 'Hủy bỏ gửi' notification detected.")

            # Final safety sleep to allow OWA to finalize SMTP handover
            print("[Playwright Raw Fallback] Sleeping 5s before closing browser...")
            await asyncio.sleep(5)
            print("[Playwright Raw Fallback] Wait complete. Closing browser.")


        except Exception as exc:
            print(f"[Playwright Raw Fallback] Error occurred during raw flow: {exc}")
            try:
                os.makedirs("C:/Users/Truon/.gemini/antigravity/brain/93fb4859-d7ac-4272-81c9-dfbf8c16bb7c/scratch/modern_survey", exist_ok=True)
                screenshot_path = "C:/Users/Truon/.gemini/antigravity/brain/93fb4859-d7ac-4272-81c9-dfbf8c16bb7c/scratch/modern_survey/error_screenshot.png"
                await page.screenshot(path=screenshot_path)
                print(f"[Playwright Raw Fallback] Captured error screenshot to {screenshot_path}")
            except Exception as ss_exc:
                print(f"[Playwright Raw Fallback] Failed to capture error screenshot: {ss_exc}")
            raise exc
        finally:
            await context.close()

    return case_data.get("to") or case_data.get("from") or "Webmail Playwright Raw Fallback"


async def send_professional_reply_all_via_playwright_raw(case_data: dict, html_body: str, professional_email: str) -> str:
    """Uses raw Playwright commands to reply-all to the email and add a professional recipient."""
    # Extract valid HTML payload to avoid pasting <html> and <head> tags into OWA's contenteditable editor.
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_body, "html.parser")
        main_table = soup.find("table")
        if main_table:
            container = soup.find(class_="container")
            if container:
                html_body = str(container)
            else:
                parts = []
                for node in main_table.previous_siblings:
                    parts.insert(0, str(node))
                parts.append(str(main_table))
                html_body = "".join(parts)
        else:
            body = soup.find("body")
            if body:
                html_body = "".join(str(c) for c in body.contents)
    except Exception as e:
        print(f"[Playwright Pro ReplyAll] Failed to extract valid HTML snippet: {e}")

    import os
    import asyncio
    from playwright.async_api import async_playwright

    username = os.getenv("MAIL_USERNAME")
    password = os.getenv("MAIL_PASSWORD")
    user_data_dir = os.getenv("CHROME_USER_DATA_DIR")
    if not user_data_dir:
        # Default to a local persistent directory to avoid logon session limits and reuse session
        user_data_dir = os.path.join(os.path.expanduser("~"), ".gemini", "antigravity", "chrome_profile")

    webmail_url = os.getenv("WEBMAIL_URL", "https://owa.cengroup.vn/").strip()
    is_headless = os.getenv("PLAYWRIGHT_HEADLESS", "true" if os.name != "nt" else "false").strip().casefold() in ("true", "1", "yes")

    contract_number = str(case_data.get("contract_number") or case_data.get("contract_id") or "").strip()
    if not contract_number:
        raise ValueError("Hồ sơ không có số hợp đồng để tìm kiếm email.")

    chrome_args = ["--no-sandbox", "--disable-dev-shm-usage", "--disable-gpu"]

    print(f"[Playwright Pro ReplyAll] Starting Playwright reply-all flow for contract {contract_number}")

    user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"

    async with async_playwright() as p:
        if user_data_dir:
            os.makedirs(user_data_dir, exist_ok=True)
            context = await p.chromium.launch_persistent_context(
                user_data_dir=user_data_dir,
                headless=is_headless,
                user_agent=user_agent,
                args=chrome_args,
                permissions=["clipboard-read", "clipboard-write"],
            )
            page = context.pages[0] if context.pages else await context.new_page()
        else:
            browser = await p.chromium.launch(
                headless=is_headless,
                args=chrome_args,
            )
            context = await browser.new_context(
                permissions=["clipboard-read", "clipboard-write"]
            )
            page = await context.new_page()

        try:
            # 1. Navigate to the login page
            print(f"[Playwright Pro ReplyAll] Navigating to {webmail_url}")
            await page.goto(webmail_url, timeout=60000)
            await page.wait_for_load_state("load")

            # Check if we are already logged in
            username_selector = "input#username, input[name='username'], input[type='email'], input[type='text']"
            try:
                await page.wait_for_selector(username_selector, timeout=5000)
                needs_login = True
            except Exception:
                needs_login = False

            if needs_login:
                print("[Playwright Pro ReplyAll] Login form detected. Logging in...")
                await page.fill("input#username, input[name='username'], input[type='email']", username)
                await page.fill("input#password, input[name='password'], input[type='password']", password)
                await page.keyboard.press("Enter")
                print("[Playwright Pro ReplyAll] Submitted login form via Enter. Waiting for navigation/dashboard...")
                await page.wait_for_load_state("networkidle", timeout=15000)
                await page.wait_for_timeout(10000)
            else:
                print("[Playwright Pro ReplyAll] Already logged in or login form not visible.")
                await page.wait_for_timeout(5000)

            # Handle OWA error pages (e.g. "too many active sessions")
            for _retry in range(3):
                error_page = await page.evaluate(
                    r"""
                    () => {
                      const text = document.body ? (document.body.innerText || '') : '';
                      if (/Something went wrong|too many active sessions|Refresh the page/i.test(text)) {
                        const refreshLink = document.querySelector('a, button, .refreshPageButton, #refreshPageButton');
                        if (refreshLink) {
                          refreshLink.click();
                          return { error: true, action: 'clicked_refresh' };
                        }
                        return { error: true, action: 'detected_but_no_refresh_link' };
                      }
                      return { error: false };
                    }
                    """
                )
                if error_page and error_page.get("error"):
                    print(f"[Playwright Pro ReplyAll] OWA error page detected: {error_page}. Refreshing...")
                    if error_page.get("action") == "clicked_refresh":
                        try:
                            await page.wait_for_load_state("load", timeout=20000)
                        except Exception:
                            pass
                    else:
                        await page.reload(wait_until="load", timeout=30000)
                    await page.wait_for_timeout(10000)
                else:
                    break

            # OWA first-run setup can appear after login on a fresh browser profile.
            try:
                first_run_done = await page.evaluate(
                r"""
                    () => {
                      const text = document.body ? (document.body.innerText || '') : '';
                      const selects = Array.from(document.querySelectorAll('select'));
                      if (!selects.length || !/Outlook|m\u00fai gi\u1edd|time zone|ng\u00f4n ng\u1eef|language/i.test(text)) {
                        return { handled: false, reason: 'first-run setup not visible' };
                      }
                      const timezoneSelect = selects[selects.length - 1];
                      const options = Array.from(timezoneSelect.options || []);
                      const preferred = options.find((option) => /Bangkok|Hanoi|Ha Noi|Jakarta|\\+07|UTC\\+7|SE Asia/i.test(option.textContent || option.value));
                      const fallback = options.find((option) => option.value && !/ch\u1ecdn|select/i.test(option.textContent || ''));
                      const chosen = preferred || fallback;
                      if (chosen) {
                        timezoneSelect.value = chosen.value;
                        timezoneSelect.dispatchEvent(new Event('change', { bubbles: true }));
                      }
                      const buttons = Array.from(document.querySelectorAll('button, input[type="submit"], [role="button"], div[role="button"]'));
                      const saveButton = buttons.find((el) => /save|l\u01b0u/i.test(el.innerText || el.textContent || el.value || el.getAttribute('aria-label') || ''));
                      if (!saveButton) return { handled: false, reason: 'save button not found' };
                      saveButton.click();
                      return { handled: true, timezone: chosen ? (chosen.textContent || chosen.value) : null };
                     }
                     """
                )
                if first_run_done and first_run_done.get("handled"):
                    print(f"[Playwright Pro ReplyAll] Completed OWA first-run setup: {first_run_done}")
                    await page.wait_for_load_state("networkidle", timeout=20000)
                    await page.wait_for_timeout(5000)
            except Exception as setup_exc:
                print(f"[Playwright Pro ReplyAll] OWA first-run setup check skipped: {setup_exc}")

            # 2. Focus OWA search input via Alt+Q shortcut.
            print("[Playwright Pro ReplyAll] Focusing OWA search input via Alt+Q...")
            await page.keyboard.press("Alt+q")
            await page.wait_for_timeout(2000)

            # Verify search input is focused
            search_focus = None
            for _ in range(10):
                search_focus = await page.evaluate(
                    r"""
                    () => {
                      const isVisible = (el) => {
                        const rect = el.getBoundingClientRect();
                        return rect.width > 0 && rect.height > 0;
                      };
                      const inputs = Array.from(document.querySelectorAll('input')).filter(isVisible);
                      const searchInput = inputs.find(i => {
                        const h = [i.type, i.title, i.placeholder, i.getAttribute('aria-label'), i.id].join(' ').toLowerCase();
                        return h.includes('search') || h.includes('tìm kiếm');
                      }) || inputs.find(i => i.type === 'text');
                      if (searchInput) {
                        searchInput.focus();
                        searchInput.click();
                        const rect = searchInput.getBoundingClientRect();
                        return { ok: true, rect: { x: rect.x, y: rect.y, w: rect.width, h: rect.height } };
                      }
                      return { ok: false };
                    }
                    """
                )
                if search_focus and search_focus.get("ok"):
                    break
                await page.wait_for_timeout(1000)

            if not search_focus or not search_focus.get("ok"):
                raise RuntimeError(f"OWA search input not found after Alt+Q: {search_focus}")
            print(f"[Playwright Pro ReplyAll] Search input focused: {search_focus}")

            # 3. Input contract number and run search.
            search_query = contract_number
            parts = [p.strip() for p in contract_number.split('/') if p.strip()]
            for p in parts:
                if '-' in p and len(p) >= 5:
                    search_query = p
                    break
            else:
                for p in parts:
                    if not p.isdigit() and p.upper() != 'DN' and len(p) >= 4:
                        search_query = p
                        break

            print(f"[Playwright Pro ReplyAll] Inputting search query: {search_query} (derived from contract: {contract_number})")
            await page.keyboard.press("Control+A")
            await page.keyboard.press("Delete")
            await page.keyboard.type(search_query, delay=20)
            search_clicked = await page.evaluate(
                r"""
                () => {
                  const isVisible = (el) => {
                    const rect = el.getBoundingClientRect();
                    const style = window.getComputedStyle(el);
                    return rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none';
                  };
                  const normalize = (value) => (value || '').toLocaleLowerCase('vi-VN');
                  const buttons = Array.from(document.querySelectorAll('button, a, [role="button"], input[type="submit"]'));
                  const searchButton = buttons.find((el) => {
                    if (!isVisible(el)) return false;
                    const text = normalize([el.title, el.getAttribute('aria-label'), el.innerText, el.textContent, el.value].filter(Boolean).join(' '));
                    const rect = el.getBoundingClientRect();
                    return (text.includes('search') || text.includes('tìm kiếm')) && rect.y < 80;
                  });
                  if (!searchButton) return false;
                  searchButton.click();
                  return true;
                }
                """
            )
            if search_clicked:
                print("[Playwright Pro ReplyAll] Search submitted via search button.")
            else:
                await page.keyboard.press("Enter")
                print("[Playwright Pro ReplyAll] Search submitted via Enter.")

            print("[Playwright Pro ReplyAll] Waiting search results to load...")
            await page.wait_for_timeout(8000)

            opened_result = await page.evaluate(
                r"""
                (contractNumber) => {
                  const normalize = (value) => (value || '').replace(/\s+/g, ' ').trim();
                  const contract = normalize(contractNumber);
                  const parts = contract.split('/');
                  const hyphenPart = parts.find(p => p.includes('-') && p.length >= 5) || '';
                  const tail = parts.filter(Boolean).pop() || contract;
                  const tailPrefix = tail.slice(0, Math.max(5, Math.min(tail.length, 7)));
                  const terms = Array.from(new Set([contract, hyphenPart, tail, tailPrefix].filter((term) => term && term.length >= 3)));
                  const isVisible = (el) => {
                    const rect = el.getBoundingClientRect();
                    const style = window.getComputedStyle(el);
                    return rect.width > 30 && rect.height > 10 && style.visibility !== 'hidden' && style.display !== 'none';
                  };
                  const clickTargetFor = (el) => el.closest('a, [role="option"], [role="listitem"], [role="button"], button, tr') || el;
                  const nodes = Array.from(document.querySelectorAll('tr, a, td, div, span'));
                  const mapped = nodes
                    .map((el) => ({ el, target: clickTargetFor(el), text: normalize(el.innerText || el.textContent), rect: el.getBoundingClientRect(), tag: el.tagName }))
                    .filter((item) => isVisible(item.el));
                  const rowLike = (item) => item.rect.y > 120 && item.rect.height >= 12 && item.rect.height < 80 && item.rect.width > 80;
                  const textMatches = mapped
                    .filter((item) => rowLike(item) && terms.some((term) => item.text.includes(term)))
                    .sort((a, b) => {
                      const aScore = (a.text.includes(contract) ? 0 : a.text.includes(hyphenPart) ? 1 : a.text.includes(tail) ? 2 : 3) + (a.tag === 'A' ? 0 : a.tag === 'TR' ? 1 : 2);
                      const bScore = (b.text.includes(contract) ? 0 : b.text.includes(hyphenPart) ? 1 : b.text.includes(tail) ? 2 : 3) + (b.tag === 'A' ? 0 : b.tag === 'TR' ? 1 : 2);
                      return (aScore - bScore) || (a.rect.y - b.rect.y) || (a.rect.x - b.rect.x);
                    });
                  const item = textMatches[0];
                  if (!item) return { ok: false, reason: 'no matching search result found', terms };
                  item.target.click();
                  return { ok: true, text: item.text.slice(0, 160), rect: { x: item.rect.x, y: item.rect.y, w: item.rect.width, h: item.rect.height }, tag: item.tag, terms };
                }
                """,
                contract_number,
            )
            if opened_result and opened_result.get("ok"):
                print(f"[Playwright Pro ReplyAll] Opened matching search result: {opened_result}")
            else:
                raise RuntimeError(f"Could not find or open matching search result: {opened_result}")

            await page.wait_for_timeout(4000)

            # 4. Open Reply All composer. Classic OWA supports a direct ReplyAll action URL.
            classic_reply_opened = False
            try:
                current_url = page.url
                if "ae=Item" in current_url and "id=" in current_url:
                    from urllib.parse import parse_qs, urlencode, urlparse, urlunparse
                    parsed = urlparse(current_url)
                    query = parse_qs(parsed.query)
                    message_id = (query.get("id") or [None])[0]
                    if message_id:
                        reply_query = {
                            "ae": "PreFormAction",
                            "a": "ReplyAll",
                            "t": query.get("t", ["IPM.Note"])[0],
                            "id": message_id,
                        }
                        reply_url = urlunparse((parsed.scheme, parsed.netloc, parsed.path, "", urlencode(reply_query), ""))
                        print(f"[Playwright Pro ReplyAll] Opening classic OWA ReplyAll URL: {reply_url}")
                        await page.goto(reply_url, wait_until="domcontentloaded", timeout=60000)
                        await page.wait_for_timeout(5000)
                        classic_reply_opened = True
            except Exception as reply_url_exc:
                print(f"[Playwright Pro ReplyAll] Classic ReplyAll URL fallback skipped: {reply_url_exc}")

            if not classic_reply_opened:
                print("[Playwright Pro ReplyAll] Locating Reply All or Continue Editing button...")
                reply_all_result = await page.evaluate(
                    r"""
                    () => {
                      const normalize = (value) => (value || '').toLocaleLowerCase('vi-VN').normalize('NFD').replace(/[\u0300-\u036f]/g, '').replace(/\\s+/g, ' ').trim();
                      const isVisible = (el) => {
                        const rect = el.getBoundingClientRect();
                        const style = window.getComputedStyle(el);
                        return rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none';
                      };
                      
                      const buttons = Array.from(document.querySelectorAll('button, a, [role="button"], [role="menuitem"]')).filter(isVisible);
                      
                      // 1. Check if "Tiếp tục chỉnh sửa" / "Continue editing" is visible
                      const continueBtn = buttons.find(el => {
                        const text = normalize([el.getAttribute('aria-label'), el.title, el.innerText, el.textContent].filter(Boolean).join(' '));
                        return text.includes('tiep tuc chinh sua') || text.includes('continue editing');
                      });
                      
                      if (continueBtn) {
                        const target = continueBtn.closest('button, [role="button"]') || continueBtn;
                        target.click();
                        return { ok: true, action: 'clicked_continue_editing', text: 'Continue Editing' };
                      }
                      
                      // 2. Fallback to Reply All button
                      const replyAllTerms = [
                        'reply all',
                        'tra loi tat ca',
                        'tra loi tat',
                        'tr\u1ea3 l\u1eddi t\u1ea5t c\u1ea3',
                        'tr\u1ea3 l\u1eddi t\u1ea5t'
                      ];
                      const hasReplyAll = (text) => {
                        const value = normalize(text);
                        return replyAllTerms.some((term) => value.includes(term));
                      };
                      
                      const candidates = buttons
                        .map((el) => ({
                          el,
                          text: [el.getAttribute('aria-label'), el.getAttribute('title'), el.innerText, el.textContent].filter(Boolean).join(' '),
                          rect: el.getBoundingClientRect(),
                        }))
                        .filter((item) => hasReplyAll(item.text))
                        .sort((a, b) => (a.rect.y - b.rect.y) || (a.rect.x - b.rect.x));
                        
                      const item = candidates[0];
                      if (!item) return { ok: false, reason: 'neither continue editing nor reply all button found' };
                      const target = item.el.closest('button, [role="button"]') || item.el;
                      target.click();
                      return { ok: true, action: 'clicked_reply_all', text: item.text.slice(0, 160) };
                    }
                    """
                )
                reply_all_clicked = bool(reply_all_result and reply_all_result.get("ok"))
                if reply_all_clicked:
                    print(f"[Playwright Pro ReplyAll] Reply All clicked via JS: {reply_all_result}")
                else:
                    print(f"[Playwright Pro ReplyAll] Reply All JS lookup failed: {reply_all_result}")
                    for selector in (
                        "button:has-text('Reply All')",
                        "button:has-text('Tr\u1ea3 l\u1eddi t\u1ea5t c\u1ea3')",
                        "[aria-label*='Reply all']",
                        "[aria-label*='Reply All']",
                        "[aria-label*='Tr\u1ea3 l\u1eddi t\u1ea5t']",
                        "[title*='Reply all']",
                        "[title*='Reply All']",
                        "[title*='Tr\u1ea3 l\u1eddi t\u1ea5t']",
                        "div[role='button']:has-text('Reply All')",
                        "span:has-text('Reply All')",
                    ):
                        try:
                            target = page.locator(selector).first
                            if await target.count() > 0:
                                await target.click(timeout=5000)
                                reply_all_clicked = True
                                print(f"[Playwright Pro ReplyAll] Reply All clicked via selector: {selector}")
                                break
                        except Exception:
                            continue

                if not reply_all_clicked:
                    print("[Playwright Pro ReplyAll] Reply All button not found. Trying Ctrl+Shift+R shortcut...")
                    await page.keyboard.press("Control+Shift+R")
                    await page.wait_for_timeout(3000)
                    reply_editor_opened = await page.evaluate(
                    r"""
                        () => Array.from(document.querySelectorAll('[contenteditable="true"], [role="textbox"], [aria-label*="Message body"], [aria-label*="body"]'))
                          .some((el) => {
                            const rect = el.getBoundingClientRect();
                            const style = window.getComputedStyle(el);
                            return rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none';
                          })
                        """
                    )
                    if not reply_editor_opened:
                        raise RuntimeError("Reply All button not found in OWA and Ctrl+Shift+R did not open reply editor.")
                    print("[Playwright Pro ReplyAll] Reply editor opened via Ctrl+Shift+R fallback.")

            await page.wait_for_timeout(3000)

            # Wait for editor to appear
            print("[Playwright Pro ReplyAll] Waiting for visible reply editor...")
            editor_found = False
            for _ in range(30):
                editor_found = await page.evaluate(
                    r"""
                    () => {
                      const isVisible = (el) => {
                        const rect = el.getBoundingClientRect();
                        const style = window.getComputedStyle(el);
                        return rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none';
                      };
                      return Array.from(document.querySelectorAll("iframe, textarea, [contenteditable='true'], [role='textbox'], [aria-label*='Message body'], [aria-label*='body']")).some(isVisible);
                    }
                    """
                )
                if editor_found:
                    break
                await page.wait_for_timeout(1000)
            if not editor_found:
                raise RuntimeError("Reply editor did not appear (no visible matching elements found).")

            # 5. Add professional recipient in "To" field
            print("[Playwright Pro ReplyAll] Locating To field to add professional recipient...")
            to_input_result = await page.evaluate(
                r"""
                () => {
                  const isVisible = (el) => {
                    const rect = el.getBoundingClientRect();
                    const style = window.getComputedStyle(el);
                    return rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none';
                  };
                  
                  const getElements = (doc) => {
                    return Array.from(doc.querySelectorAll("input, textarea, div"));
                  };
                  
                  const isToField = (i) => {
                    const label = [i.getAttribute('aria-label'), i.title, i.placeholder, i.id, i.name].filter(Boolean).join(' ').toLowerCase();
                    const className = i.className || '';
                    
                    const matchesText = (label === 'to' || label === 'tới' || label.includes('recipient') || label.includes('to:') || label.includes('tới:') || label.includes('đến người nhận')) &&
                                        !label.includes('subject') && !label.includes('cc') && !label.includes('body') && !label.includes('chủ đề');
                    
                    const isTargetTag = i.tagName === 'INPUT' || i.tagName === 'TEXTAREA' || i.contentEditable === 'true' || i.getAttribute('role') === 'textbox' || className.includes('_rw_') || className.includes('RecipientWell');
                    
                    return matchesText && isTargetTag;
                  };

                  let inputs = getElements(document).filter(isVisible);
                  const iframes = Array.from(document.querySelectorAll('iframe')).filter(isVisible);
                  let frameOffset = { x: 0, y: 0 };
                  let foundInFrame = false;
                  let toInput = null;

                  for (const frame of iframes) {
                    try {
                      const doc = frame.contentDocument || frame.contentWindow.document;
                      if (doc) {
                        const frameInputs = getElements(doc).filter(isVisible);
                        const match = frameInputs.find(isToField);
                        if (match) {
                          toInput = match;
                          const frameRect = frame.getBoundingClientRect();
                          frameOffset = { x: frameRect.x, y: frameRect.y };
                          foundInFrame = true;
                          break;
                        }
                      }
                    } catch (e) {}
                  }

                  if (!toInput) {
                    toInput = inputs.find(isToField);
                  }

                  if (toInput) {
                    toInput.focus();
                    toInput.click();
                    const rect = toInput.getBoundingClientRect();
                    return { 
                      ok: true, 
                      rect: { 
                        x: frameOffset.x + rect.x + rect.width / 2, 
                        y: frameOffset.y + rect.y + rect.height / 2 
                      } 
                    };
                  }
                  return { ok: false, reason: 'To input not found in main document or iframes' };
                }
                """
            )
            if to_input_result and to_input_result.get("ok"):
                rect = to_input_result["rect"]
                if rect["x"] > 0 and rect["y"] > 0:
                    print(f"[Playwright Pro ReplyAll] Clicking To input field at: {rect}")
                    await page.mouse.click(rect["x"], rect["y"])
                    await page.wait_for_timeout(500)
                else:
                    print("[Playwright Pro ReplyAll] To field coordinates are 0,0. Relying on JS focus and click.")
                print(f"[Playwright Pro ReplyAll] Typing professional email: {professional_email}")
                await page.keyboard.type(professional_email, delay=20)
                await page.wait_for_timeout(500)
                await page.keyboard.press("Semicolon")
                await page.keyboard.press("Enter")
                print("[Playwright Pro ReplyAll] Recipient added.")
                await page.wait_for_timeout(1000)
            else:
                print(f"[Playwright Pro ReplyAll] Warning: Could not locate To input to add recipient: {to_input_result.get('reason')}")
                try:
                    diag_elements = await page.evaluate(
                        r"""
                        () => {
                          const isVisible = (el) => {
                            const rect = el.getBoundingClientRect();
                            const style = window.getComputedStyle(el);
                            return rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none';
                          };
                          const getElData = (doc, prefix = '') => {
                            return Array.from(doc.querySelectorAll('input, textarea, div[contenteditable="true"], div[role="textbox"]'))
                              .filter(isVisible)
                              .map(el => ({
                                tag: el.tagName,
                                id: el.id,
                                className: el.className,
                                ariaLabel: el.getAttribute('aria-label'),
                                title: el.title,
                                placeholder: el.placeholder || el.getAttribute('placeholder'),
                                role: el.getAttribute('role'),
                                text: el.innerText ? el.innerText.slice(0, 100) : '',
                                rect: el.getBoundingClientRect()
                              }));
                          };
                          let res = getElData(document).map(x => ({...x, frame: 'main'}));
                          const frames = Array.from(document.querySelectorAll('iframe')).filter(isVisible);
                          frames.forEach((frame, idx) => {
                            try {
                              const doc = frame.contentDocument || frame.contentWindow.document;
                              if (doc) {
                                res = res.concat(getElData(doc).map(x => ({...x, frame: `iframe_${idx}`})));
                              }
                            } catch(e) {}
                          });
                          return res;
                        }
                        """
                    )
                    print(f"[Playwright Pro ReplyAll] Diagnostic: Found {len(diag_elements)} visible inputs/textboxes:")
                    for idx, el in enumerate(diag_elements):
                        print(f"  - {idx}: Frame={el['frame']}, Tag={el['tag']}, ID={el['id']}, Class={el['className']}, AriaLabel={el['ariaLabel']}, Title={el['title']}, Placeholder={el['placeholder']}, Text='{el['text']}'")
                    
                    # Target composer header buttons
                    diag_buttons = await page.evaluate(
                        r"""
                        () => {
                          const isVisible = (el) => {
                            const rect = el.getBoundingClientRect();
                            const style = window.getComputedStyle(el);
                            return rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none';
                          };
                          const editor = document.querySelector('.owa-font-compose');
                          if (!editor) return [{ error: 'editor not found' }];
                          
                          let composer = editor;
                          while (composer && composer !== document.body) {
                            if (composer.className && (composer.className.includes('composer') || composer.className.includes('Compose') || composer.className.includes('ReadingPane') || composer.className.includes('_rp_'))) {
                              break;
                            }
                            composer = composer.parentElement;
                          }
                          if (!composer) composer = document;
                          
                          return Array.from(composer.querySelectorAll('button, a, [role="button"], span, div'))
                            .filter(isVisible)
                            .filter(el => {
                              const text = (el.innerText || el.textContent || '').trim();
                              const title = el.title || '';
                              const aria = el.getAttribute('aria-label') || '';
                              const className = el.className || '';
                              const isClickable = el.tagName === 'BUTTON' || el.tagName === 'A' || el.getAttribute('role') === 'button' || className.includes('button') || className.includes('Btn');
                              const matchesTerms = /pop|cửa sổ|expand|tới|to|cc|bcc|người nhận|recipient|chỉnh sửa|edit/i.test(text + ' ' + title + ' ' + aria + ' ' + className);
                              return isClickable || (matchesTerms && (text.length < 50 || title || aria));
                            })
                            .map(el => ({
                              tag: el.tagName,
                              id: el.id,
                              className: el.className,
                              text: el.innerText ? el.innerText.slice(0, 50).replace(/\n/g, ' ') : '',
                              ariaLabel: el.getAttribute('aria-label'),
                              title: el.title,
                              rect: el.getBoundingClientRect()
                            }));
                        }
                        """
                    )
                    print(f"[Playwright Pro ReplyAll] Diagnostic: Found {len(diag_buttons)} composer header buttons:")
                    for idx, btn in enumerate(diag_buttons):
                        if 'error' in btn:
                            print(f"  - Error: {btn['error']}")
                        else:
                            print(f"  - {idx}: Tag={btn['tag']}, Class={btn['className']}, Text='{btn['text']}', AriaLabel={btn['ariaLabel']}, Title={btn['title']}, Rect={btn['rect']}")
                except Exception as diag_exc:
                    print(f"[Playwright Pro ReplyAll] Diagnostic dump failed: {diag_exc}")

            # 6. Locate, click, and insert html_body into editor.
            print("[Playwright Pro ReplyAll] Locating editor coordinates for native click...")
            editor_loc = await page.evaluate(
                r"""
                () => {
                  const isVisible = (el) => {
                    const rect = el.getBoundingClientRect();
                    const style = window.getComputedStyle(el);
                    return rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none';
                  };

                  let editor = null;
                  let frameOffset = { x: 0, y: 0 };
                  let mode = 'contenteditable';

                  // 1. Check iframes
                  const iframes = Array.from(document.querySelectorAll('iframe')).filter(isVisible);
                  for (const frame of iframes) {
                    try {
                      const doc = frame.contentDocument || frame.contentWindow.document;
                      if (doc && doc.body && doc.body.getAttribute('contenteditable') === 'true') {
                        editor = doc.body;
                        const frameRect = frame.getBoundingClientRect();
                        frameOffset = { x: frameRect.x, y: frameRect.y };
                        mode = 'iframe';
                        break;
                      }
                    } catch (e) {}
                  }

                  // 2. Check main document for owa-font-compose
                  if (!editor) {
                    editor = document.querySelector('div.owa-font-compose[contenteditable="true"]');
                  }

                  // 3. Candidates fallback with strict exclusions
                  if (!editor) {
                    const candidates = Array.from(document.querySelectorAll('[contenteditable="true"], [role="textbox"]')).filter(isVisible);
                    editor = candidates.find(c => {
                      const label = (c.getAttribute('aria-label') || '').toLowerCase();
                      const className = typeof c.className === 'string' ? c.className.toLowerCase() : '';
                      if (label.includes('to') || label.includes('tới') || label.includes('recipient') || 
                          label.includes('cc') || label.includes('subject') || label.includes('chủ đề') ||
                          label.includes('tìm kiếm') || label.includes('search')) {
                        return false;
                      }
                      if (className.includes('well') || className.includes('_rw_') || className.includes('recipient') || className.includes('search')) {
                        return false;
                      }
                      return true;
                    });
                  }

                  // 4. Textarea fallback with strict exclusions
                  if (!editor) {
                    const textareas = Array.from(document.querySelectorAll('textarea')).filter(isVisible);
                    const bodyTextarea = textareas.find(t => {
                      const label = (t.getAttribute('aria-label') || '').toLowerCase();
                      const className = typeof t.className === 'string' ? t.className.toLowerCase() : '';
                      if (label.includes('to') || label.includes('tới') || label.includes('recipient') || 
                          label.includes('cc') || label.includes('subject') || label.includes('chủ đề')) {
                        return false;
                      }
                      if (className.includes('well') || className.includes('_rw_') || className.includes('recipient')) {
                        return false;
                      }
                      return true;
                    });
                    if (bodyTextarea) {
                      editor = bodyTextarea;
                      mode = 'textarea';
                    }
                  }

                  if (editor) {
                    try {
                      editor.scrollIntoView({ block: 'nearest', inline: 'nearest' });
                    } catch (e) {}
                    const rect = editor.getBoundingClientRect();
                    return {
                      ok: true,
                      mode,
                      x: frameOffset.x + rect.x + 50,
                      y: frameOffset.y + rect.y + 20
                    };
                  }
                  return { ok: false };
                }
                """
            )
            print(f"[Playwright Pro ReplyAll] Editor location result: {editor_loc}")
            if not editor_loc or not editor_loc.get("ok"):
                raise RuntimeError("Could not locate OWA compose body editor.")

            print(f"[Playwright Pro ReplyAll] Performing native click on editor at ({editor_loc['x']}, {editor_loc['y']})")
            await page.mouse.click(editor_loc["x"], editor_loc["y"])
            await page.wait_for_timeout(300)

            # 6. Insert html_body into editor.
            print("[Playwright Pro ReplyAll] Writing HTML body to clipboard...")
            plain_body = html_to_plain_text(html_body)
            clipboard_result = await page.evaluate(
                r"""
                async (args) => {
                  try {
                    if (typeof ClipboardItem === 'undefined') {
                      return { ok: false, error: 'ClipboardItem is not supported' };
                    }
                    const htmlBlob = new Blob([args.html], { type: 'text/html' });
                    const plainBlob = new Blob([args.plain], { type: 'text/plain' });
                    const item = new ClipboardItem({
                      'text/html': htmlBlob,
                      'text/plain': plainBlob
                    });
                    await navigator.clipboard.write([item]);
                    return { ok: true };
                  } catch (e) {
                    return { ok: false, error: e.toString() };
                  }
                }
                """,
                {"html": html_body, "plain": plain_body}
            )
            print(f"[Playwright Pro ReplyAll] Clipboard write result: {clipboard_result}")

            pasted_ok = False
            if clipboard_result and clipboard_result.get("ok"):
                print("[Playwright Pro ReplyAll] Pasting HTML content natively using Control+v...")
                await page.keyboard.press("Control+v")
                await page.wait_for_timeout(1000)
                
                # Check if rich HTML paste worked (must contain table tag)
                pasted_ok = await page.evaluate(
                    r"""
                    () => {
                      let activeDoc = document;
                      if (document.activeElement && document.activeElement.tagName === 'IFRAME') {
                        try {
                          activeDoc = document.activeElement.contentDocument || document.activeElement.contentWindow.document;
                        } catch (e) {}
                      }
                      const editor = activeDoc.activeElement;
                      if (!editor) return false;
                      const html = editor.innerHTML || '';
                      return html.includes('<table') || html.includes('<TABLE');
                    }
                    """
                )
                print(f"[Playwright Pro ReplyAll] Native paste verification (has table): {pasted_ok}")

            # Fallback to execCommand / innerHTML + space/backspace keystroke if clipboard paste failed
            if not pasted_ok:
                print("[Playwright Pro ReplyAll] Native paste failed. Falling back to execCommand/innerHTML...")
                result = await page.evaluate(
                    r"""
                    (args) => {
                      const html = args.html;
                      const plain = args.plain;
                      const mode = args.mode;
                      
                      let activeDoc = document;
                      if (document.activeElement && document.activeElement.tagName === 'IFRAME') {
                        try {
                          activeDoc = document.activeElement.contentDocument || document.activeElement.contentWindow.document;
                        } catch (e) {}
                      }
                      
                      const editor = activeDoc.activeElement;
                      if (!editor) {
                        return { ok: false, reason: 'No active element focused in document' };
                      }

                      if (mode === 'textarea' || editor.tagName === 'TEXTAREA') {
                        editor.value = plain;
                        editor.dispatchEvent(new Event('input', { bubbles: true }));
                        editor.dispatchEvent(new Event('change', { bubbles: true }));
                        return { ok: true, mode: 'textarea' };
                      }

                      let execOk = false;
                      try {
                        const win = activeDoc.defaultView || window;
                        win.focus();
                        editor.focus();
                        const sel = win.getSelection();
                        sel.removeAllRanges();
                        const range = activeDoc.createRange();
                        range.selectNodeContents(editor);
                        range.collapse(true); // collapse to beginning of text
                        sel.addRange(range);
                        execOk = activeDoc.execCommand('insertHTML', false, html);
                      } catch (e) {
                        console.error(e);
                      }

                      if (!execOk) {
                        editor.innerHTML = html + editor.innerHTML;
                      }

                      editor.dispatchEvent(new Event('input', { bubbles: true }));
                      editor.dispatchEvent(new Event('change', { bubbles: true }));
                      
                      const finalHTML = editor.innerHTML || '';
                      return {
                        ok: finalHTML.includes('<table') || finalHTML.includes('<TABLE'),
                        mode: 'contenteditable',
                        execOk
                      };
                    }
                    """,
                    {"html": html_body, "plain": plain_body, "mode": editor_loc["mode"]},
                )
                print(f"[Playwright Pro ReplyAll] Editor insertion fallback result: {result}")
                
                print(f"[Playwright Pro ReplyAll] Re-clicking editor at ({editor_loc['x']}, {editor_loc['y']}) to ensure cursor active...")
                await page.mouse.click(editor_loc["x"], editor_loc["y"])
                await page.wait_for_timeout(200)
                
                # Programmatically place selection cursor at the start of editor contents
                await page.evaluate(
                    r"""
                    () => {
                      let activeDoc = document;
                      if (document.activeElement && document.activeElement.tagName === 'IFRAME') {
                        try {
                          activeDoc = document.activeElement.contentDocument || document.activeElement.contentWindow.document;
                        } catch (e) {}
                      }
                      const editor = activeDoc.activeElement;
                      if (!editor) return;
                      const win = activeDoc.defaultView || window;
                      win.focus();
                      editor.focus();
                      const sel = win.getSelection();
                      sel.removeAllRanges();
                      const range = activeDoc.createRange();
                      range.selectNodeContents(editor);
                      range.collapse(true);
                      sel.addRange(range);
                    }
                    """
                )
                await page.wait_for_timeout(100)
                
                print("[Playwright Pro ReplyAll] Simulating native keystroke (Space + Backspace) to trigger editor state sync...")
                await page.keyboard.press("Space")
                await page.wait_for_timeout(200)
                await page.keyboard.press("Backspace")
                await page.wait_for_timeout(500)

            # 7. Send email via Ctrl+Enter
            print("[Playwright Pro ReplyAll] Sending email via Ctrl+Enter...")
            await page.evaluate(
                r"""
                () => {
                  let activeDoc = document;
                  if (document.activeElement && document.activeElement.tagName === 'IFRAME') {
                    try {
                      activeDoc = document.activeElement.contentDocument || document.activeElement.contentWindow.document;
                    } catch (e) {}
                  }
                  let editor = activeDoc.querySelector('div.owa-font-compose[contenteditable="true"]');
                  if (!editor && activeDoc.body && activeDoc.body.getAttribute('contenteditable') === 'true') {
                    editor = activeDoc.body;
                  }
                  if (!editor) {
                    const candidates = Array.from(activeDoc.querySelectorAll('[contenteditable="true"]')).filter(el => {
                      const r = el.getBoundingClientRect();
                      return r.width > 200 && r.height > 50;
                    });
                    if (candidates.length) {
                      editor = candidates[candidates.length - 1];
                    }
                  }
                  if (editor) {
                    editor.focus();
                  }
                }
                """
            )
            await page.wait_for_timeout(500)
            await page.keyboard.press("Control+Enter")
            print("[Playwright Pro ReplyAll] Ctrl+Enter sent.")

            # 8. Wait for compose editor to close + Fallback Send button click
            print("[Playwright Pro ReplyAll] Waiting for editor to close...")
            editor_closed = False
            try:
                await page.wait_for_selector('div.owa-font-compose[contenteditable="true"]', state="hidden", timeout=3000)
                editor_closed = True
                print("[Playwright Pro ReplyAll] Editor closed.")
            except Exception:
                print("[Playwright Pro ReplyAll] Editor did not close within 3s. Triggering fallback Send button click...")
                # Attempt to click "Gửi" / "Send" button programmatically
                btn_clicked = await page.evaluate(
                    r"""
                    () => {
                      const normalize = (val) => (val || '').trim().toLowerCase();
                      const sendBtn = Array.from(document.querySelectorAll('button, a, [role="button"]')).find(el => {
                        const text = normalize(el.innerText || el.textContent || '');
                        const ariaLabel = normalize(el.getAttribute('aria-label') || '');
                        const title = normalize(el.getAttribute('title') || '');
                        return text === 'gửi' || text === 'send' || ariaLabel === 'gửi' || ariaLabel === 'send' || title === 'gửi' || title === 'send';
                      });
                      if (sendBtn) {
                        const target = sendBtn.closest('button, [role="button"]') || sendBtn;
                        target.click();
                        return true;
                      }
                      return false;
                    }
                    """
                )
                print(f"[Playwright Pro ReplyAll] Programmatic Send button clicked: {btn_clicked}")
            
            if not editor_closed:
                try:
                    await page.wait_for_selector('div.owa-font-compose[contenteditable="true"]', state="hidden", timeout=7000)
                    print("[Playwright Pro ReplyAll] Editor closed after fallback click.")
                except Exception:
                    print("[Playwright Pro ReplyAll] Editor did not close within 10s total.")

            # Wait for "Hủy bỏ gửi"
            print("[Playwright Pro ReplyAll] Checking for OWA 'Hủy bỏ gửi' progress bar...")
            has_undo = False
            for _ in range(6):
                undo_exists = await page.evaluate(
                    r"""
                    () => {
                      const normalize = (val) => (val || '').trim().toLowerCase();
                      const els = Array.from(document.querySelectorAll('span, div, button, a'));
                      return els.some(el => {
                        const text = normalize(el.innerText || el.textContent || '');
                        return text.includes('hủy bỏ gửi') || text.includes('hủy gửi') || text.includes('undo send');
                      });
                    }
                    """
                )
                if undo_exists:
                    has_undo = True
                    print("[Playwright Pro ReplyAll] 'Hủy bỏ gửi' progress bar detected. Waiting for it to complete...")
                    break
                await page.wait_for_timeout(500)

            if has_undo:
                for _ in range(60):
                    undo_exists = await page.evaluate(
                        r"""
                        () => {
                          const normalize = (val) => (val || '').trim().toLowerCase();
                          const els = Array.from(document.querySelectorAll('span, div, button, a'));
                          return els.some(el => {
                            const text = normalize(el.innerText || el.textContent || '');
                            const rect = el.getBoundingClientRect();
                            return (text.includes('hủy bỏ gửi') || text.includes('hủy gửi') || text.includes('undo send')) && rect.width > 0 && rect.height > 0;
                          });
                        }
                        """
                    )
                    if not undo_exists:
                        print("[Playwright Pro ReplyAll] 'Hủy bỏ gửi' progress bar completed. Mail is sent!")
                        break
                    await page.wait_for_timeout(500)
            else:
                print("[Playwright Pro ReplyAll] No 'Hủy bỏ gửi' notification detected.")

            print("[Playwright Pro ReplyAll] Sleeping 5s before closing browser...")
            await asyncio.sleep(5)
            print("[Playwright Pro ReplyAll] Wait complete. Closing browser.")

        except Exception as exc:
            print(f"[Playwright Pro ReplyAll] Error occurred during raw flow: {exc}")
            try:
                screenshot_dir = os.path.join(os.getcwd(), "logs")
                os.makedirs(screenshot_dir, exist_ok=True)
                screenshot_path = os.path.join(screenshot_dir, "error_professional_forward.png")
                await page.screenshot(path=screenshot_path)
                print(f"[Playwright Pro ReplyAll] Captured error screenshot to {screenshot_path}")
            except Exception as ss_exc:
                print(f"[Playwright Pro ReplyAll] Failed to capture error screenshot: {ss_exc}")
            raise exc
        finally:
            await context.close()

    return professional_email


