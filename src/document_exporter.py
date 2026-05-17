from __future__ import annotations

import copy
import re
import html
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .case_files import case_folder, sanitize_folder_name
from .database_store import format_money, money_to_vietnamese_words, parse_money
from .template_manager import PLACEHOLDER_PATTERN, iter_paragraphs, read_docx_text


# --- Các trường tài sản có thể chứa nhiều dòng (đa tài sản) ---
MULTILINE_CONTEXT_KEYS = {
    "TAI_SAN_THAM_DINH",
    "DIA_CHI_TAI_SAN",
}


def _expand_multiline_paragraph(paragraph: Paragraph) -> None:
    """Nếu paragraph chứa ký tự xuống dòng, tách thành nhiều paragraph con
    với gạch đầu dòng (- ). Giữ nguyên format (font, size) của run gốc."""
    full_text = paragraph.text
    if "\n" not in full_text:
        return

    lines = [line.strip() for line in full_text.split("\n") if line.strip()]
    if len(lines) <= 1:
        return

    # Lấy element gốc và parent
    p_element = paragraph._element
    parent = p_element.getparent()
    if parent is None:
        return

    # Lưu lại thông tin format từ run đầu tiên
    ref_run = paragraph.runs[0] if paragraph.runs else None

    # Đặt nội dung paragraph hiện tại = dòng đầu tiên (có gạch đầu dòng)
    first_line = f"- {lines[0]}"
    if paragraph.runs:
        paragraph.runs[0].text = first_line
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(first_line)

    # Tạo các paragraph mới cho các dòng còn lại, chèn ngay sau paragraph hiện tại
    insert_after = p_element
    for line in lines[1:]:
        new_p = copy.deepcopy(p_element)
        # Xóa hết run cũ trong bản sao
        for r in new_p.findall(qn("w:r")):
            new_p.remove(r)
        # Tạo run mới
        new_run = copy.deepcopy(ref_run._element) if ref_run is not None else None
        if new_run is not None:
            # Xóa text cũ trong run
            for t in new_run.findall(qn("w:t")):
                new_run.remove(t)
            from docx.oxml import OxmlElement
            t_elem = OxmlElement("w:t")
            t_elem.set(qn("xml:space"), "preserve")
            t_elem.text = f"- {line}"
            new_run.append(t_elem)
            new_p.append(new_run)
        else:
            from docx.oxml import OxmlElement
            new_run_elem = OxmlElement("w:r")
            t_elem = OxmlElement("w:t")
            t_elem.set(qn("xml:space"), "preserve")
            t_elem.text = f"- {line}"
            new_run_elem.append(t_elem)
            new_p.append(new_run_elem)
        insert_after.addnext(new_p)
        insert_after = new_p


CONTRACT_DOC_SUFFIX = "/HĐTĐG"
PAYMENT_REQUEST_SUFFIX = "/TT-CENVALUE-ĐN"

INDIVIDUAL_TEMPLATES = {
    "hop_dong": ("mau_hd.docx", "Hop dong", "{contract}_hop_dong.docx"),
    "phieu_yeu_cau": ("mau_pyc.docx", "Phieu yeu cau", "{contract}_phieu_yeu_cau.docx"),
    "bien_ban_nghiem_thu": (
        "mau_bbnt.docx",
        "Bien ban nghiem thu",
        "{contract}_bien_ban_nghiem_thu.docx",
    ),
}

ORGANIZATION_TEMPLATES = {
    "hop_dong": ("hop_dong_vcb.docx", "Hop dong to chuc", "{contract}_hop_dong_to_chuc.docx"),
    "bien_ban_nghiem_thu": (
        "bbtl_cong_ty.docx",
        "Bien ban nghiem thu/to chuc",
        "{contract}_bien_ban_nghiem_thu_to_chuc.docx",
    ),
    "de_nghi_thanh_toan": (
        "de_nghi_thanh_toan.docx",
        "De nghi thanh toan",
        "{contract}_de_nghi_thanh_toan.docx",
    ),
    "thu_chao_phi": ("thu_chao_phi.docx", "Thu chao phi", "{contract}_thu_chao_phi.docx"),
}


def clean_customer_name(customer_info: str) -> str:
    text = customer_info or ""
    text = re.sub(r"\([^)]*\)", "", text)
    text = re.sub(r"\b\d{9,12}\b", "", text)
    text = re.sub(r"\s*[-–]\s*$", "", text)
    text = re.sub(r"\s+", " ", text).strip(" -–")
    return text


def extract_phone(customer_info: str) -> str:
    match = re.search(r"\b(0\d{8,10})\b", customer_info or "")
    return match.group(1) if match else ""


def ensure_period(value: str) -> str:
    text = (value or "").strip()
    if not text:
        return ""
    return text if text.endswith((".", "!", "?")) else f"{text}."


BANK_ALIASES = {
    "vcb": "vietcombank",
    "vietcombank": "vietcombank",
    "sacombank": "sacombank",
    "vp bank": "vpbank",
    "vpbank": "vpbank",
    "techcombank": "techcombank",
    "tcb": "techcombank",
    "bidv": "bidv",
    "mb bank": "mbbank",
    "mbbank": "mbbank",
    "mb": "mbbank",
    "agribank": "agribank",
    "vib": "vib",
}


def _normalized_bank_terms(value: str) -> set[str]:
    text = re.sub(r"[^0-9a-zA-ZÀ-ỹ]+", " ", value or "").casefold()
    compact = text.replace(" ", "")
    terms: set[str] = set()
    for alias, canonical in BANK_ALIASES.items():
        alias_text = alias.casefold()
        if alias_text in text or alias_text.replace(" ", "") in compact:
            terms.add(canonical)
    return terms


def _should_append_source_to_purpose(purpose: str, source: str) -> bool:
    if not source:
        return False
    if "ngân hàng" in purpose.casefold():
        return False
    if source.casefold() in purpose.casefold():
        return False
    if re.fullmatch(r"[A-ZĐ]{2,5}", source.strip()):
        return False
    source_bank_terms = _normalized_bank_terms(source)
    if source_bank_terms and source_bank_terms.intersection(_normalized_bank_terms(purpose)):
        return False
    return True


def _protect_nonbreaking_terms(value: str) -> str:
    text = value or ""
    replacements = {
        "VP Bank": "VP\u00a0Bank",
        "MB Bank": "MB\u00a0Bank",
    }
    for old, new in replacements.items():
        text = re.sub(re.escape(old), new, text, flags=re.IGNORECASE)
    return text


def _shorten_valuation_purpose(value: str) -> str:
    text = (value or "").strip()
    prefixes = [
        r"làm\s+cơ\s+sở\s+tham\s+khảo\s+để\s+",
        r"làm\s+cơ\s+sở\s+để\s+",
    ]
    for pattern in prefixes:
        text = re.sub(f"^{pattern}", "", text, flags=re.IGNORECASE).strip()
    return _protect_nonbreaking_terms(text)


def _extract_land_address_from_asset(value: str) -> str:
    text = (value or "").strip().rstrip(".")
    if not text:
        return ""
    match = re.search(r"\btại\s+địa\s+chỉ\s+(.+)$", text, flags=re.IGNORECASE)
    if match:
        return match.group(1).strip().rstrip(".")
    return ""


def contract_document_number(contract_number: str) -> str:
    text = (contract_number or "").strip()
    if not text:
        return ""
    return text if text.endswith(CONTRACT_DOC_SUFFIX) else f"{text}{CONTRACT_DOC_SUFFIX}"


def _date_context(date: datetime | None = None) -> dict[str, str]:
    value = date or datetime.now()
    return {
        "NGAY": f"{value.day:02d}",
        "THANG": f"{value.month:02d}",
        "NAM": str(value.year),
        "THANG_NAM": f"Tháng {value.month:02d} năm {value.year}",
        "NGAY_LAP_PLEIKU": f"Pleiku, Ngày {value.day:02d} tháng {value.month:02d} năm {value.year}",
    }


def _date_parts_from_text(value: Any) -> dict[str, str]:
    text = str(value or "").strip()
    if not text:
        return {"full": "", "day": "", "month": "", "year": ""}
    match = re.search(r"(\d{1,2})[./-](\d{1,2})[./-](\d{4})", text)
    if match:
        day, month, year = match.groups()
        return {
            "full": f"{int(day):02d}/{int(month):02d}/{year}",
            "day": f"{int(day):02d}",
            "month": f"{int(month):02d}",
            "year": year,
        }
    match = re.search(r"(\d{4})[./-](\d{1,2})[./-](\d{1,2})", text)
    if match:
        year, month, day = match.groups()
        return {
            "full": f"{int(day):02d}/{int(month):02d}/{year}",
            "day": f"{int(day):02d}",
            "month": f"{int(month):02d}",
            "year": year,
        }
    return {"full": text, "day": "", "month": "", "year": ""}


def _pleiku_date_from_parts(parts: dict[str, str]) -> str:
    if parts["day"] and parts["month"] and parts["year"]:
        return f"Pleiku, Ngày {parts['day']} tháng {parts['month']} năm {parts['year']}"
    return _date_context()["NGAY_LAP_PLEIKU"]


def _compute_remaining_amount(case: dict[str, Any]) -> str:
    fee = parse_money(case.get("valuation_fee_number"))
    advance = parse_money(case.get("advance_payment"))
    if fee is None:
        return ""
    if advance is None:
        advance = 0
    return format_money(max(fee - advance, 0))


def _fee_words(case: dict[str, Any]) -> str:
    explicit = str(case.get("valuation_fee_words") or "").strip()
    if explicit and not explicit.startswith("#"):
        return explicit.rstrip(".")
    fee = parse_money(case.get("valuation_fee_number"))
    if fee is None:
        return ""
    return money_to_vietnamese_words(fee).rstrip(".")


def build_placeholder_context(case: dict[str, Any], *, organization: bool = False) -> dict[str, str]:
    customer_info = str(case.get("customer_info") or "").strip()
    customer_name = clean_customer_name(customer_info)
    phone = str(case.get("customer_phone") or "").strip()
    if not phone:
        phone = extract_phone(customer_info)
    contract_number = str(case.get("contract_number") or "").strip()
    source = str(case.get("source") or "").strip()
    purpose = str(case.get("valuation_purpose") or "").strip()
    purpose_full = purpose
    if _should_append_source_to_purpose(purpose, source):
        purpose_full = f"{purpose} {source}".strip()
    purpose = _protect_nonbreaking_terms(purpose)
    purpose_full = _protect_nonbreaking_terms(purpose_full)
    contract_date = _date_parts_from_text(case.get("contract_date"))
    certificate_date = _date_parts_from_text(case.get("certificate_date"))
    asset_description = str(case.get("asset_description") or "").strip()
    land_address = str(
        case.get("dia_chi_thua_dat")
        or _extract_land_address_from_asset(asset_description)
        or case.get("customer_address")
        or ""
    ).strip()

    context = {
        "TEN_KHACH_HANG": customer_name,
        "DIA_CHI_KHACH_HANG": str(case.get("customer_address") or "").strip(),
        "CCCD": str(case.get("citizen_id") or "").strip(),
        "DIEN_THOAI_KHACH_HANG": phone,
        "TAI_SAN_THAM_DINH": ensure_period(asset_description),
        "DIA_CHI_TAI_SAN": land_address,
        "MUC_DICH_THAM_DINH": purpose,
        "MUC_DICH_THAM_DINH_DAY_DU": ensure_period(purpose_full),
        "MUC_DICH_THAM_DINH_RUT_GON": ensure_period(_shorten_valuation_purpose(purpose)),
        "NGUON": source,
        "SO_HOP_DONG": contract_number,
        "SO_HOP_DONG_VAN_BAN": contract_document_number(contract_number),
        "NGAY_HOP_DONG": contract_date["full"],
        "NGAY_HOP_DONG_PLEIKU": _pleiku_date_from_parts(contract_date),
        "NGAY_HOP_DONG_NGAY": contract_date["day"],
        "NGAY_HOP_DONG_THANG": contract_date["month"],
        "NGAY_HOP_DONG_NAM": contract_date["year"],
        "NGAY_CHUNG_THU": certificate_date["full"],
        "NGAY_CHUNG_THU_NGAY": certificate_date["day"],
        "NGAY_CHUNG_THU_THANG": certificate_date["month"],
        "NGAY_CHUNG_THU_NAM": certificate_date["year"],
        "SO_BIEN_BAN_NGHIEM_THU": f"{contract_number}/BBNT" if contract_number else "",
        "SO_DE_NGHI_THANH_TOAN": f"{contract_number}{PAYMENT_REQUEST_SUFFIX}" if contract_number else "",
        "PHI_THAM_DINH": format_money(case.get("valuation_fee_number")),
        "PHI_THAM_DINH_BANG_CHU": _fee_words(case),
        "TAM_UNG": format_money(case.get("advance_payment")),
        "CON_LAI_THANH_TOAN": _compute_remaining_amount(case),
    }
    context.update(_date_context())

    if organization:
        context.update(
            {
                "MA_SO_THUE": str(case.get("tax_code") or "").strip(),
                "NGUOI_DAI_DIEN": str(case.get("representative_name") or "").strip(),
                "CHUC_VU_NGUOI_DAI_DIEN": str(case.get("representative_position") or "").strip(),
                "CAN_CU_UY_QUYEN": str(case.get("authorization_note") or "").strip(),
                "NGUOI_NHAN_BAN_GIAO": str(case.get("handover_contact_name") or "").strip(),
                "CHUC_VU_NGUOI_NHAN_BAN_GIAO": str(case.get("handover_contact_position") or "").strip(),
                "SDT_NGUOI_NHAN_BAN_GIAO": str(case.get("handover_contact_phone") or "").strip(),
            }
        )

    return context


def _replace_placeholders(text: str, context: dict[str, str]) -> str:
    updated = text
    for key, value in sorted(context.items(), key=lambda item: len(item[0]), reverse=True):
        updated = updated.replace(f"{{{{{key}}}}}", value or "")
    return updated


def _replace_in_paragraph(paragraph, context: dict[str, str]) -> None:
    original = paragraph.text
    updated = _replace_placeholders(original, context)
    if updated == original:
        return
    if not paragraph.runs:
        paragraph.add_run(updated)
        return

    run_texts = [run.text for run in paragraph.runs]
    full_text = "".join(run_texts)
    if full_text != original:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
        return

    ranges: list[tuple[int, int]] = []
    cursor = 0
    for text in run_texts:
        start = cursor
        cursor += len(text)
        ranges.append((start, cursor))

    replacements = [
        (match.start(), match.end(), str(context.get(match.group(1)) or ""))
        for match in PLACEHOLDER_PATTERN.finditer(full_text)
        if match.group(1) in context
    ]
    if not replacements:
        return

    for start, end, replacement in reversed(replacements):
        affected = [
            index
            for index, (run_start, run_end) in enumerate(ranges)
            if run_start < end and run_end > start
        ]
        if not affected:
            continue
        first_index = affected[0]
        for index in affected:
            run_start, run_end = ranges[index]
            local_start = max(start, run_start) - run_start
            local_end = min(end, run_end) - run_start
            text = run_texts[index]
            insert = replacement if index == first_index else ""
            run_texts[index] = text[:local_start] + insert + text[local_end:]

    for run, text in zip(paragraph.runs, run_texts):
        run.text = text


def _alignment_css(value: Any) -> str:
    if value is None:
        return ""
    alignment = int(value)
    if alignment == 1:
        return "text-align:center;"
    if alignment == 2:
        return "text-align:right;"
    if alignment == 3:
        return "text-align:justify;"
    return "text-align:left;"


def _paragraph_style_css(paragraph: Paragraph) -> str:
    style_name = (paragraph.style.name if paragraph.style else "").lower()
    rules = [_alignment_css(paragraph.alignment)]
    if "heading 1" in style_name or "tiêu đề 1" in style_name:
        rules.append("font-size:20pt;font-weight:700;margin:14pt 0 8pt;")
    elif "heading 2" in style_name or "tiêu đề 2" in style_name:
        rules.append("font-size:16pt;font-weight:700;margin:12pt 0 6pt;")
    elif "title" in style_name or "tiêu đề" in style_name:
        rules.append("font-size:18pt;font-weight:700;text-align:center;margin:12pt 0;")
    return "".join(rule for rule in rules if rule)


def _run_style_css(run) -> str:
    rules: list[str] = []
    if run.bold:
        rules.append("font-weight:700;")
    if run.italic:
        rules.append("font-style:italic;")
    if run.underline:
        rules.append("text-decoration:underline;")
    if run.font.size:
        rules.append(f"font-size:{run.font.size.pt:.1f}pt;")
    if run.font.name:
        rules.append(f"font-family:{html.escape(run.font.name)}, 'Times New Roman', serif;")
    if run.font.color and run.font.color.rgb:
        rules.append(f"color:#{run.font.color.rgb};")
    return "".join(rules)


def _render_paragraph_html(paragraph: Paragraph, context: dict[str, str]) -> str:
    original = paragraph.text
    rendered = _replace_placeholders(original, context)
    paragraph_style = _paragraph_style_css(paragraph)
    if not rendered.strip():
        return f'<p style="{paragraph_style}">&nbsp;</p>'

    # Nếu kết quả chứa xuống dòng (đa tài sản), render dưới dạng danh sách gạch đầu dòng
    if "\n" in rendered:
        lines = [line.strip() for line in rendered.split("\n") if line.strip()]
        if len(lines) > 1:
            run_style = _run_style_css(paragraph.runs[0]) if paragraph.runs else ""
            items = "".join(f'<li style="{run_style}">{html.escape(line)}</li>' for line in lines)
            return f'<ul style="{paragraph_style}margin-left:18px;padding-left:0;">{items}</ul>'

    parts: list[str] = []
    rendered_run_texts: list[str] = []
    for run in paragraph.runs:
        text = _replace_placeholders(run.text, context)
        rendered_run_texts.append(text)
        if text:
            parts.append(f'<span style="{_run_style_css(run)}">{html.escape(text)}</span>')
    run_rendered = "".join(rendered_run_texts)
    if parts and "{{" not in run_rendered and "}}" not in run_rendered:
        body = "".join(parts)
    elif rendered != original:
        run_style = _run_style_css(paragraph.runs[0]) if paragraph.runs else ""
        body = f'<span style="{run_style}">{html.escape(rendered)}</span>'
    else:
        body = "".join(parts) or html.escape(rendered)
    return f'<p style="{paragraph_style}">{body}</p>'


def _render_table_html(table: Table, context: dict[str, str]) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            content = "".join(_render_paragraph_html(paragraph, context) for paragraph in cell.paragraphs)
            cells.append(f"<td>{content}</td>")
        rows.append(f"<tr>{''.join(cells)}</tr>")
    return f"<table>{''.join(rows)}</table>"


def _iter_block_items(document):
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def render_docx_preview_html(
    template_path: str | Path,
    case: dict[str, Any],
    *,
    organization: bool = False,
) -> str:
    document = Document(str(template_path))
    context = build_placeholder_context(case, organization=organization)
    blocks: list[str] = []
    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            blocks.append(_render_paragraph_html(block, context))
        elif isinstance(block, Table):
            blocks.append(_render_table_html(block, context))
    return f"""
    <style>
      .word-preview {{
        box-sizing: border-box;
        width: 794px;
        min-height: 1123px;
        margin: 0 auto;
        padding: 72px 76px;
        background: #fff;
        color: #111827;
        font-family: "Times New Roman", serif;
        font-size: 12pt;
        line-height: 1.35;
        box-shadow: 0 8px 30px rgba(15, 23, 42, 0.15);
      }}
      .word-preview p {{
        margin: 0 0 7pt;
        white-space: pre-wrap;
      }}
      .word-preview table {{
        width: 100%;
        border-collapse: collapse;
        margin: 8pt 0;
      }}
      .word-preview td {{
        border: 1px solid #222;
        padding: 4pt 6pt;
        vertical-align: top;
      }}
      body {{
        margin: 0;
        background: #f1f5f9;
        padding: 18px;
      }}
    </style>
    <div class="word-preview">{''.join(blocks)}</div>
    """


def _case_folder_and_contract(case: dict[str, Any], case_files_dir: str | Path) -> tuple[Path, str]:
    case_id = int(case["id"])
    folder = Path(
        case.get("case_folder")
        or case_folder(
            case_files_dir,
            case_id=case_id,
            contract_number=case.get("contract_number") or "",
            customer_name=case.get("customer_info") or "",
        )
    )
    contract = sanitize_folder_name(case.get("contract_number") or f"HS-{case_id:05d}", fallback=f"HS-{case_id:05d}")
    return folder, contract


def _describe_documents(
    templates: dict[str, tuple[str, str, str]],
    *,
    case: dict[str, Any],
    templates_dir: str | Path,
    case_files_dir: str | Path,
) -> list[dict[str, Any]]:
    folder, contract = _case_folder_and_contract(case, case_files_dir)
    output_dir = folder
    documents: list[dict[str, Any]] = []
    for key, (template_name, label, output_pattern) in templates.items():
        documents.append(
            {
                "key": key,
                "name": label,
                "template": Path(templates_dir) / template_name,
                "output_path": output_dir / output_pattern.format(contract=contract),
            }
        )
    return documents


def describe_individual_documents(
    case: dict[str, Any],
    *,
    templates_dir: str | Path,
    case_files_dir: str | Path,
) -> list[dict[str, Any]]:
    return _describe_documents(
        INDIVIDUAL_TEMPLATES,
        case=case,
        templates_dir=templates_dir,
        case_files_dir=case_files_dir,
    )


def describe_organization_documents(
    case: dict[str, Any],
    *,
    templates_dir: str | Path,
    case_files_dir: str | Path,
) -> list[dict[str, Any]]:
    return _describe_documents(
        ORGANIZATION_TEMPLATES,
        case=case,
        templates_dir=templates_dir,
        case_files_dir=case_files_dir,
    )


def render_docx_template(
    template_path: str | Path,
    output_path: str | Path,
    case: dict[str, Any],
    *,
    organization: bool = False,
) -> Path:
    document = Document(str(template_path))
    context = build_placeholder_context(case, organization=organization)
    # Thu thập tất cả paragraph trước để tránh vấn đề khi thêm paragraph mới
    all_paragraphs = list(iter_paragraphs(document))
    for paragraph in all_paragraphs:
        _replace_in_paragraph(paragraph, context)

    # Sau khi thay thế, tách các paragraph chứa \n thành danh sách gạch đầu dòng
    for paragraph in all_paragraphs:
        _expand_multiline_paragraph(paragraph)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def render_docx_preview(
    template_path: str | Path,
    case: dict[str, Any],
    *,
    organization: bool = False,
) -> str:
    document = Document(str(template_path))
    context = build_placeholder_context(case, organization=organization)
    lines: list[str] = []
    for paragraph in iter_paragraphs(document):
        rendered = _replace_placeholders(paragraph.text.strip(), context).strip()
        if rendered:
            lines.append(rendered)
    return "\n".join(lines)


def _preview_from_descriptions(
    descriptions: list[dict[str, Any]],
    case: dict[str, Any],
    *,
    organization: bool,
) -> list[dict[str, str]]:
    previews: list[dict[str, str]] = []
    for item in descriptions:
        template = Path(item["template"])
        if not template.exists():
            raise FileNotFoundError(template)
        previews.append(
            {
                "key": item["key"],
                "name": item["name"],
                "template": str(template),
                "output_path": str(item["output_path"]),
                "content": render_docx_preview(template, case, organization=organization),
                "html": render_docx_preview_html(template, case, organization=organization),
            }
        )
    return previews


def preview_individual_document_set(
    case: dict[str, Any],
    *,
    templates_dir: str | Path,
    case_files_dir: str | Path,
) -> list[dict[str, str]]:
    descriptions = describe_individual_documents(case, templates_dir=templates_dir, case_files_dir=case_files_dir)
    return _preview_from_descriptions(descriptions, case, organization=False)


def preview_organization_document_set(
    case: dict[str, Any],
    *,
    templates_dir: str | Path,
    case_files_dir: str | Path,
) -> list[dict[str, str]]:
    descriptions = describe_organization_documents(case, templates_dir=templates_dir, case_files_dir=case_files_dir)
    return _preview_from_descriptions(descriptions, case, organization=True)


def _export_from_descriptions(
    descriptions: list[dict[str, Any]],
    case: dict[str, Any],
    *,
    organization: bool,
) -> list[Path]:
    output_paths: list[Path] = []
    for item in descriptions:
        template = Path(item["template"])
        if not template.exists():
            raise FileNotFoundError(template)
        output_paths.append(render_docx_template(template, item["output_path"], case, organization=organization))
    return output_paths


def export_individual_document_set(
    case: dict[str, Any],
    *,
    templates_dir: str | Path,
    case_files_dir: str | Path,
) -> list[Path]:
    descriptions = describe_individual_documents(case, templates_dir=templates_dir, case_files_dir=case_files_dir)
    return _export_from_descriptions(descriptions, case, organization=False)


def export_organization_document_set(
    case: dict[str, Any],
    *,
    templates_dir: str | Path,
    case_files_dir: str | Path,
) -> list[Path]:
    descriptions = describe_organization_documents(case, templates_dir=templates_dir, case_files_dir=case_files_dir)
    return _export_from_descriptions(descriptions, case, organization=True)


def compare_preview_with_export(previews: list[dict[str, str]]) -> list[dict[str, Any]]:
    comparisons: list[dict[str, Any]] = []
    for preview in previews:
        output_path = Path(preview["output_path"])
        if not output_path.exists():
            comparisons.append(
                {
                    "name": preview["name"],
                    "output_path": str(output_path),
                    "matched": False,
                    "reason": "Chua tim thay file Word da xuat.",
                    "exported_content": "",
                    "preview_content": preview["content"],
                }
            )
            continue

        exported_content = read_docx_text(output_path)
        matched = exported_content == preview["content"]
        comparisons.append(
            {
                "name": preview["name"],
                "output_path": str(output_path),
                "matched": matched,
                "reason": "" if matched else "Noi dung preview va file da xuat dang khac nhau.",
                "exported_content": exported_content,
                "preview_content": preview["content"],
            }
        )
    return comparisons
