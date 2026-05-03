from __future__ import annotations

import json
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

from docx import Document


PLACEHOLDER_PATTERN = re.compile(r"\{\{([A-Z0-9_]+)\}\}")
TEMPLATE_REQUIREMENTS = {
    "individual": {
        "mau_hd.docx": {
            "TEN_KHACH_HANG",
            "DIA_CHI_KHACH_HANG",
            "CCCD",
            "DIEN_THOAI_KHACH_HANG",
            "TAI_SAN_THAM_DINH",
            "MUC_DICH_THAM_DINH_DAY_DU",
            "SO_HOP_DONG_VAN_BAN",
            "PHI_THAM_DINH",
            "PHI_THAM_DINH_BANG_CHU",
            "NGAY",
            "THANG",
            "NAM",
        },
        "mau_pyc.docx": {
            "TEN_KHACH_HANG",
            "DIA_CHI_KHACH_HANG",
            "DIEN_THOAI_KHACH_HANG",
            "CCCD",
            "TAI_SAN_THAM_DINH",
            "DIA_CHI_TAI_SAN",
            "MUC_DICH_THAM_DINH_DAY_DU",
            "MUC_DICH_THAM_DINH_RUT_GON",
            "NGAY_HOP_DONG_PLEIKU",
            "THANG_NAM",
        },
        "mau_bbnt.docx": {
            "TEN_KHACH_HANG",
            "DIA_CHI_KHACH_HANG",
            "CCCD",
            "SO_HOP_DONG",
            "SO_HOP_DONG_VAN_BAN",
            "NGAY",
            "THANG",
            "NAM",
        },
    },
    "organization": {
        "hop_dong_vcb.docx": {
            "TEN_KHACH_HANG",
            "DIA_CHI_KHACH_HANG",
            "DIEN_THOAI_KHACH_HANG",
            "MA_SO_THUE",
            "NGUOI_DAI_DIEN",
            "CHUC_VU_NGUOI_DAI_DIEN",
            "TAI_SAN_THAM_DINH",
            "MUC_DICH_THAM_DINH_DAY_DU",
            "SO_HOP_DONG_VAN_BAN",
            "PHI_THAM_DINH",
            "NGAY",
            "THANG",
            "NAM",
        },
        "bbtl_cong_ty.docx": {
            "TEN_KHACH_HANG",
            "DIA_CHI_KHACH_HANG",
            "MA_SO_THUE",
            "NGUOI_DAI_DIEN",
            "CHUC_VU_NGUOI_DAI_DIEN",
            "SO_BIEN_BAN_NGHIEM_THU",
            "SO_HOP_DONG_VAN_BAN",
            "PHI_THAM_DINH",
            "PHI_THAM_DINH_BANG_CHU",
            "TAM_UNG",
            "NGAY",
            "THANG",
            "NAM",
        },
        "de_nghi_thanh_toan.docx": {
            "TEN_KHACH_HANG",
            "DIA_CHI_KHACH_HANG",
            "SO_DE_NGHI_THANH_TOAN",
            "SO_HOP_DONG_VAN_BAN",
            "PHI_THAM_DINH",
            "TAM_UNG",
            "CON_LAI_THANH_TOAN",
            "NGAY",
            "THANG",
            "NAM",
        },
        "thu_chao_phi.docx": {
            "TEN_KHACH_HANG",
            "TAI_SAN_THAM_DINH",
            "MUC_DICH_THAM_DINH_DAY_DU",
            "PHI_THAM_DINH",
            "PHI_THAM_DINH_BANG_CHU",
            "NGAY",
            "THANG",
            "NAM",
        },
    },
}


def iter_paragraph_nodes(document: Document):
    for index, paragraph in enumerate(document.paragraphs, start=1):
        yield {
            "paragraph": paragraph,
            "block_id": f"body-p-{index}",
            "location": f"Body P{index}",
            "scope": "body",
        }

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for col_index, cell in enumerate(row.cells, start=1):
                for para_index, paragraph in enumerate(cell.paragraphs, start=1):
                    yield {
                        "paragraph": paragraph,
                        "block_id": f"table-{table_index}-r-{row_index}-c-{col_index}-p-{para_index}",
                        "location": f"Table {table_index} / R{row_index} C{col_index} / P{para_index}",
                        "scope": "table",
                    }

    for section_index, section in enumerate(document.sections, start=1):
        for para_index, paragraph in enumerate(section.header.paragraphs, start=1):
            yield {
                "paragraph": paragraph,
                "block_id": f"header-{section_index}-p-{para_index}",
                "location": f"Header {section_index} / P{para_index}",
                "scope": "header",
            }
        for para_index, paragraph in enumerate(section.footer.paragraphs, start=1):
            yield {
                "paragraph": paragraph,
                "block_id": f"footer-{section_index}-p-{para_index}",
                "location": f"Footer {section_index} / P{para_index}",
                "scope": "footer",
            }


def iter_paragraphs(document: Document):
    for node in iter_paragraph_nodes(document):
        yield node["paragraph"]


def _set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def normalize_template_path(path: str | Path) -> str:
    return str(Path(path).resolve()).lower()


def get_template_label(path: str | Path, labels: dict[str, str]) -> str:
    return str(labels.get(normalize_template_path(path), "draft")).strip() or "draft"


def set_template_label(path: str | Path, labels: dict[str, str], label: str) -> dict[str, str]:
    updated = dict(labels)
    updated[normalize_template_path(path)] = label.strip() or "draft"
    return updated


def load_template_config(config_path: str | Path, defaults: dict[str, Any]) -> dict[str, Any]:
    path = Path(config_path)
    if not path.exists():
        return dict(defaults)

    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return dict(defaults)

    merged = dict(defaults)
    for key, value in data.items():
        if key not in merged:
            continue
        if isinstance(merged[key], list) and isinstance(value, list):
            merged[key] = [str(item) for item in value]
        elif isinstance(merged[key], dict) and isinstance(value, dict):
            merged[key] = {str(k): str(v) for k, v in value.items()}
        elif isinstance(merged[key], str) and isinstance(value, str):
            merged[key] = value
    return merged


def save_template_config(config_path: str | Path, config: dict[str, Any]) -> None:
    path = Path(config_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    payload: dict[str, Any] = {}
    for key, value in config.items():
        if isinstance(value, list):
            payload[key] = [str(item) for item in value]
        elif isinstance(value, dict):
            payload[key] = {str(k): str(v) for k, v in value.items()}
        else:
            payload[key] = str(value)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def extract_placeholders_from_docx(path: str | Path) -> list[str]:
    document = Document(str(path))
    placeholders: set[str] = set()
    for paragraph in iter_paragraphs(document):
        placeholders.update(PLACEHOLDER_PATTERN.findall(paragraph.text))
    return sorted(placeholders)


def list_editable_blocks(path: str | Path, *, placeholders_only: bool = False) -> list[dict[str, Any]]:
    document = Document(str(path))
    blocks: list[dict[str, Any]] = []
    for node in iter_paragraph_nodes(document):
        paragraph = node["paragraph"]
        text = paragraph.text.strip()
        placeholders = sorted(set(PLACEHOLDER_PATTERN.findall(paragraph.text)))
        if not text:
            continue
        if placeholders_only and not placeholders:
            continue
        blocks.append(
            {
                "block_id": node["block_id"],
                "location": node["location"],
                "scope": node["scope"],
                "text": paragraph.text,
                "placeholders": placeholders,
            }
        )
    return blocks


def update_template_blocks(path: str | Path, updates: dict[str, str]) -> list[dict[str, str]]:
    document = Document(str(path))
    changes: list[dict[str, str]] = []
    for node in iter_paragraph_nodes(document):
        block_id = node["block_id"]
        if block_id not in updates:
            continue
        new_text = updates[block_id]
        paragraph = node["paragraph"]
        if paragraph.text == new_text:
            continue
        old_text = paragraph.text
        _set_paragraph_text(paragraph, new_text)
        changes.append(
            {
                "block_id": block_id,
                "location": str(node["location"]),
                "old_text": old_text,
                "new_text": new_text,
            }
        )

    if changes:
        document.save(str(path))
    return changes


def read_docx_text(path: str | Path) -> str:
    document = Document(str(path))
    lines: list[str] = []
    for paragraph in iter_paragraphs(document):
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    return "\n".join(lines)


def required_placeholders_for_template(path: str | Path, customer_type: str) -> list[str]:
    template_name = Path(path).name.lower()
    required = TEMPLATE_REQUIREMENTS.get(customer_type, {}).get(template_name, set())
    return sorted(required)


def validate_template_placeholders(path: str | Path, customer_type: str) -> dict[str, list[str]]:
    present = extract_placeholders_from_docx(path)
    required = required_placeholders_for_template(path, customer_type)
    present_set = set(present)
    required_set = set(required)
    missing = sorted(required_set - present_set)
    extra = sorted(present_set - required_set) if required else []
    return {
        "present": present,
        "required": required,
        "missing": missing,
        "extra": extra,
    }


def list_docx_templates(directory: str | Path) -> list[Path]:
    path = Path(directory)
    if not path.exists() or not path.is_dir():
        return []
    return sorted(path.glob("*.docx"))


def is_template_locked(path: str | Path, locked_templates: list[str]) -> bool:
    normalized = normalize_template_path(path)
    return normalized in {normalize_template_path(item) for item in locked_templates}


def set_template_lock(path: str | Path, locked_templates: list[str], *, locked: bool) -> list[str]:
    normalized = normalize_template_path(path)
    items = {normalize_template_path(item) for item in locked_templates}
    if locked:
        items.add(normalized)
    else:
        items.discard(normalized)
    return sorted(items)


def append_template_history(
    history_path: str | Path,
    *,
    template_path: str | Path,
    editor_name: str,
    action: str,
    details: dict[str, Any],
) -> None:
    path = Path(history_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "template_path": str(Path(template_path).resolve()),
        "template_name": Path(template_path).name,
        "editor_name": editor_name.strip() or "Unknown",
        "action": action,
        "details": details,
    }
    with path.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(payload, ensure_ascii=False) + "\n")


def create_template_snapshot(
    template_path: str | Path,
    versions_root: str | Path,
    *,
    reason: str,
) -> Path:
    source = Path(template_path)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    destination_dir = Path(versions_root) / source.stem
    destination_dir.mkdir(parents=True, exist_ok=True)
    destination = destination_dir / f"{timestamp}_{reason}{source.suffix}"
    shutil.copy2(source, destination)
    return destination


def read_template_history(
    history_path: str | Path,
    *,
    template_path: str | Path | None = None,
    limit: int = 20,
) -> list[dict[str, Any]]:
    path = Path(history_path)
    if not path.exists():
        return []

    target = normalize_template_path(template_path) if template_path else None
    rows: list[dict[str, Any]] = []
    with path.open("r", encoding="utf-8") as handle:
        for line in handle:
            line = line.strip()
            if not line:
                continue
            try:
                item = json.loads(line)
            except json.JSONDecodeError:
                continue
            if target and normalize_template_path(item.get("template_path", "")) != target:
                continue
            rows.append(item)

    return list(reversed(rows[-limit:]))


def restore_template_from_snapshot(
    template_path: str | Path,
    snapshot_path: str | Path,
) -> None:
    source = Path(snapshot_path)
    target = Path(template_path)
    if not source.exists():
        raise FileNotFoundError(source)
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, target)


def replace_text_in_docx(path: str | Path, replacements: dict[str, str]) -> bool:
    document = Document(str(path))
    changed = False
    for paragraph in iter_paragraphs(document):
        updated = paragraph.text
        for old, new in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
            updated = updated.replace(old, new)

        if updated != paragraph.text:
            changed = True
            _set_paragraph_text(paragraph, updated)

    if changed:
        document.save(str(path))
    return changed


def normalize_purpose_placeholders(paths: Iterable[str | Path]) -> list[Path]:
    replacements = {
        "{{MUC_DICH_THAM_DINH}} {{NGUON}}.": "{{MUC_DICH_THAM_DINH_DAY_DU}}",
        "{{MUC_DICH_THAM_DINH}} {{NGUON}}": "{{MUC_DICH_THAM_DINH_DAY_DU}}",
    }
    changed_paths: list[Path] = []
    for path in paths:
        target = Path(path)
        if target.exists() and replace_text_in_docx(target, replacements):
            changed_paths.append(target)
    return changed_paths
