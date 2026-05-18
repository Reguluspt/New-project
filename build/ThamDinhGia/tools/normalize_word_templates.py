from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from src.document_exporter import export_individual_document_set, export_organization_document_set
from src.template_manager import normalize_purpose_placeholders


TEMPLATES_ROOT = ROOT / "samples" / "templates"
INDIVIDUAL_DIR = TEMPLATES_ROOT / "individual"
INDIVIDUAL_RAW_DIR = TEMPLATES_ROOT / "individual_raw"
ORGANIZATION_DIR = TEMPLATES_ROOT / "organization"
WORK_DIR = ROOT / "outputs" / "template_normalization"
DUMMY_PHONE = "0900000000"


PLACEHOLDER_CASE_INDIVIDUAL = {
    "id": 1,
    "contract_number": "{{SO_HOP_DONG}}",
    "customer_type": "individual",
    "customer_info": f"{{{{TEN_KHACH_HANG}}}} - {DUMMY_PHONE}",
    "customer_address": "{{DIA_CHI_KHACH_HANG}}",
    "citizen_id": "{{CCCD}}",
    "asset_description": "{{TAI_SAN_THAM_DINH}}",
    "valuation_purpose": "{{MUC_DICH_THAM_DINH}}",
    "source": "{{NGUON}}",
    "valuation_fee_number": "{{PHI_THAM_DINH}}",
    "valuation_fee_words": "{{PHI_THAM_DINH_BANG_CHU}}",
    "dia_chi_thua_dat": "{{DIA_CHI_TAI_SAN}}",
    "case_folder": str(WORK_DIR / "individual"),
}


PLACEHOLDER_CASE_ORGANIZATION = {
    "id": 2,
    "contract_number": "{{SO_HOP_DONG}}",
    "customer_type": "organization",
    "customer_info": f"{{{{TEN_KHACH_HANG}}}} - {DUMMY_PHONE}",
    "customer_address": "{{DIA_CHI_KHACH_HANG}}",
    "tax_code": "{{MA_SO_THUE}}",
    "representative_name": "{{NGUOI_DAI_DIEN}}",
    "representative_position": "{{CHUC_VU_NGUOI_DAI_DIEN}}",
    "authorization_note": "{{CAN_CU_UY_QUYEN}}",
    "handover_contact_name": "{{NGUOI_NHAN_BAN_GIAO}}",
    "handover_contact_position": "{{CHUC_VU_NGUOI_NHAN_BAN_GIAO}}",
    "handover_contact_phone": "{{SDT_NGUOI_NHAN_BAN_GIAO}}",
    "asset_description": "{{TAI_SAN_THAM_DINH}}",
    "valuation_purpose": "{{MUC_DICH_THAM_DINH}}",
    "source": "{{NGUON}}",
    "valuation_fee_number": 9876543,
    "valuation_fee_words": "{{PHI_THAM_DINH_BANG_CHU}}",
    "advance_payment": 1234567,
    "case_folder": str(WORK_DIR / "organization"),
}


DATE_REPLACEMENTS = {
    "Pleiku, Ngày 24 tháng 04 năm 2026": "{{NGAY_LAP_PLEIKU}}",
    "Tháng 04 năm 2026": "{{THANG_NAM}}",
    "ngày  24  tháng 04  năm 2026": "ngày  {{NGAY}}  tháng {{THANG}}  năm {{NAM}}",
    "ngày  24 tháng 04 năm 2026": "ngày  {{NGAY}} tháng {{THANG}} năm {{NAM}}",
    "ngày 24 tháng 04 năm 2026": "ngày {{NGAY}} tháng {{THANG}} năm {{NAM}}",
    "TP HCM, ngày 24 tháng 04 năm 2026": "TP HCM, ngày {{NGAY}} tháng {{THANG}} năm {{NAM}}",
    "Tp. Hồ Chí Minh, ngày 24 tháng 04 năm 2026": "Tp. Hồ Chí Minh, ngày {{NGAY}} tháng {{THANG}} năm {{NAM}}",
}


GLOBAL_REPLACEMENTS = {
    "{{SO_HOP_DONG}}/HĐTĐG": "{{SO_HOP_DONG_VAN_BAN}}",
    "{{SO_HOP_DONG}}/BBNT": "{{SO_BIEN_BAN_NGHIEM_THU}}",
    "{{SO_HOP_DONG}}/TT-CENVALUE-ĐN": "{{SO_DE_NGHI_THANH_TOAN}}",
    f"{{{{TEN_KHACH_HANG}}}} Số điện thoại: {DUMMY_PHONE}": "{{TEN_KHACH_HANG}} Số điện thoại: {{DIEN_THOAI_KHACH_HANG}}",
    DUMMY_PHONE: "{{DIEN_THOAI_KHACH_HANG}}",
    "{{TEN_KHACH_HANG}} - {{DIEN_THOAI_KHACH_HANG}}": "{{TEN_KHACH_HANG}}",
    "{{MUC_DICH_THAM_DINH}} {{NGUON}}.": "{{MUC_DICH_THAM_DINH_DAY_DU}}",
    "{{MUC_DICH_THAM_DINH}} {{NGUON}}": "{{MUC_DICH_THAM_DINH_DAY_DU}}",
}


ORGANIZATION_NUMERIC_REPLACEMENTS = {
    "9.876.543": "{{PHI_THAM_DINH}}",
    "1.234.567": "{{TAM_UNG}}",
    "8.641.976": "{{CON_LAI_THANH_TOAN}}",
    "9.876.543 VNĐ": "{{PHI_THAM_DINH}} VNĐ",
    "1.234.567 VNĐ": "{{TAM_UNG}} VNĐ",
    "8.641.976 VNĐ": "{{CON_LAI_THANH_TOAN}} VNĐ",
}


def iter_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_text_in_document(path: Path, replacements: dict[str, str]) -> None:
    document = Document(str(path))

    for paragraph in iter_paragraphs(document):
        updated = paragraph.text
        for old, new in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
            updated = updated.replace(old, new)

        if path.name == "mau_pyc.docx":
            if updated.startswith("Tài sản thẩm định: "):
                updated = "Tài sản thẩm định: {{TAI_SAN_THAM_DINH}}."
            if updated.startswith("Người liên hệ khảo sát: {{TEN_KHACH_HANG}} Số điện thoại:"):
                updated = "Người liên hệ khảo sát: {{TEN_KHACH_HANG}} Số điện thoại: {{DIEN_THOAI_KHACH_HANG}}"

        if path.name == "mau_hd.docx" and updated.strip() == "Điện thoại:":
            updated = "Điện thoại: {{DIEN_THOAI_KHACH_HANG}}"

        if updated != paragraph.text:
            set_paragraph_text(paragraph, updated)

    if path.name == "de_nghi_thanh_toan.docx" and len(document.tables) >= 2:
        table = document.tables[1]
        fee_cells = [(3, 2), (3, 3), (3, 4)]
        advance_cells = [(3, 5), (3, 6)]
        remaining_cells = [(3, 7), (3, 8)]
        for row, col in fee_cells:
            for paragraph in table.cell(row, col).paragraphs:
                if paragraph.text.strip():
                    set_paragraph_text(paragraph, "{{PHI_THAM_DINH}}")
        for row, col in advance_cells:
            for paragraph in table.cell(row, col).paragraphs:
                if paragraph.text.strip():
                    set_paragraph_text(paragraph, "{{TAM_UNG}}")
        for row, col in remaining_cells:
            for paragraph in table.cell(row, col).paragraphs:
                if paragraph.text.strip():
                    set_paragraph_text(paragraph, "{{CON_LAI_THANH_TOAN}}")

    document.save(path)


def backup_individual_templates() -> None:
    INDIVIDUAL_RAW_DIR.mkdir(parents=True, exist_ok=True)
    for template in INDIVIDUAL_DIR.glob("*.docx"):
        backup = INDIVIDUAL_RAW_DIR / template.name
        if not backup.exists():
            shutil.copy2(template, backup)


def normalize_templates() -> None:
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    backup_individual_templates()

    individual_exports = export_individual_document_set(
        PLACEHOLDER_CASE_INDIVIDUAL,
        templates_dir=INDIVIDUAL_DIR,
        case_files_dir=WORK_DIR,
    )
    organization_exports = export_organization_document_set(
        PLACEHOLDER_CASE_ORGANIZATION,
        templates_dir=ORGANIZATION_DIR,
        case_files_dir=WORK_DIR,
    )

    destination_map = {
        individual_exports[0]: INDIVIDUAL_DIR / "mau_hd.docx",
        individual_exports[1]: INDIVIDUAL_DIR / "mau_pyc.docx",
        individual_exports[2]: INDIVIDUAL_DIR / "mau_bbnt.docx",
        organization_exports[0]: ORGANIZATION_DIR / "hop_dong_vcb.docx",
        organization_exports[1]: ORGANIZATION_DIR / "bbtl_cong_ty.docx",
        organization_exports[2]: ORGANIZATION_DIR / "de_nghi_thanh_toan.docx",
        organization_exports[3]: ORGANIZATION_DIR / "thu_chao_phi.docx",
    }

    for exported, destination in destination_map.items():
        shutil.copy2(exported, destination)

    replacements = dict(DATE_REPLACEMENTS)
    replacements.update(GLOBAL_REPLACEMENTS)

    for path in list(INDIVIDUAL_DIR.glob("*.docx")) + list(ORGANIZATION_DIR.glob("*.docx")):
        path_replacements = dict(replacements)
        if path.parent == ORGANIZATION_DIR:
            path_replacements.update(ORGANIZATION_NUMERIC_REPLACEMENTS)
        replace_text_in_document(path, path_replacements)

    normalize_purpose_placeholders(list(INDIVIDUAL_DIR.glob("*.docx")) + list(ORGANIZATION_DIR.glob("*.docx")))

    shutil.rmtree(WORK_DIR, ignore_errors=True)


if __name__ == "__main__":
    normalize_templates()
